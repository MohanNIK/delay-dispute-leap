from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, Iterable, List, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
IJPM_DIR = DESKTOP / "IJPM-PDW-2rd"
PAPER_ASSETS = PROJECT_ROOT / "paper_assets"

BASE_DRAFT = max((p for p in DESKTOP.glob("paper4*.docx") if p.stat().st_size > 0), key=lambda p: p.stat().st_size)
ABSTRACT_DRAFT = next(
    p
    for p in IJPM_DIR.glob("*DelayDispute Copilot*摘要.docx")
    if not p.name.startswith("~$")
)
OUTPUT_DOCX = IJPM_DIR / "DelayDispute_Copilot_Integrated_SCI_20260416.docx"


FIGURES_DIR = PAPER_ASSETS / "figures"
TABLES_DIR = PAPER_ASSETS / "tables"
TEXT_DIR = PAPER_ASSETS / "text"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 15), ("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10.5)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True


def clean_text(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\s*\[Use [^\]]+\]\s*$", "", text)
    text = text.replace("鈥?", "–").replace("鈥檚", "’s")
    return text


def parse_md_sections(path: Path) -> Dict[str, Dict[str, List[str]]]:
    raw = path.read_text(encoding="utf-8")
    sections: Dict[str, Dict[str, List[str]]] = {}
    current_top = None
    current_sub = None
    buffer: List[str] = []

    def flush():
        nonlocal buffer, current_top, current_sub
        if current_top is None or current_sub is None:
            buffer = []
            return
        text = clean_text("\n".join(buffer).strip())
        if text:
            sections.setdefault(current_top, {}).setdefault(current_sub, [])
            sections[current_top][current_sub] = [clean_text(p) for p in text.split("\n\n") if clean_text(p)]
        buffer = []

    for line in raw.splitlines():
        if line.startswith("## "):
            flush()
            current_top = line[3:].strip()
            current_sub = "__main__"
        elif line.startswith("### "):
            flush()
            current_sub = line[4:].strip()
        else:
            buffer.append(line)
    flush()
    return sections


def extract_doc_block(doc: Document, start_heading: str, end_heading: str | None = None) -> List[tuple[str, str]]:
    in_block = False
    rows: List[tuple[str, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not in_block:
            if text == start_heading:
                in_block = True
                rows.append((para.style.name, text))
            continue
        if end_heading and text == end_heading:
            break
        if text:
            rows.append((para.style.name, text))
    return rows


def add_paragraph(doc: Document, text: str, style: str = "Normal", italic: bool = False, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if style == "Normal":
        run.font.size = Pt(10.5)
    return p


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_paragraph(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def format_value(value):
    if pd.isna(value):
        return ""
    if isinstance(value, (int, float)):
        if abs(value - int(value)) < 1e-9 and abs(value) >= 1:
            return f"{int(value)}"
        return f"{value:.3f}"
    return str(value)


def style_word_table(table, font_size: float = 8.5):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(font_size)
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "space": 0, "color": "777777"},
                bottom={"val": "single", "sz": 4, "space": 0, "color": "777777"},
                left={"val": "single", "sz": 2, "space": 0, "color": "C0C0C0"},
                right={"val": "single", "sz": 2, "space": 0, "color": "C0C0C0"},
            )


def add_table_from_df(doc: Document, df: pd.DataFrame, caption: str, note: str | None = None, font_size: float = 8.5) -> None:
    add_paragraph(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = format_value(value)
    style_word_table(table, font_size=font_size)
    if note:
        add_paragraph(doc, f"Note: {note}", italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)


def compact_table_1() -> pd.DataFrame:
    df = pd.read_csv(TABLES_DIR / "table1_candidate_benchmark_profile.csv")
    return pd.DataFrame(
        {
            "Benchmark": df["benchmark"],
            "n": df["n_cases"],
            "Support": [f"{n} ({r:.3f})" for n, r in zip(df["support_n"], df["support_ratio"])],
            "Partial support": [f"{n} ({r:.3f})" for n, r in zip(df["partial_n"], df["partial_ratio"])],
            "Not support": [f"{n} ({r:.3f})" for n, r in zip(df["not_support_n"], df["not_support_ratio"])],
            "Owner": [f"{n} ({r:.3f})" for n, r in zip(df["owner_n"], df["owner_ratio"])],
            "Contractor": [f"{n} ({r:.3f})" for n, r in zip(df["contractor_n"], df["contractor_ratio"])],
            "Unknown": [f"{n} ({r:.3f})" for n, r in zip(df["unknown_n"], df["unknown_ratio"])],
        }
    )


def compact_table_2() -> pd.DataFrame:
    df = pd.read_csv(TABLES_DIR / "table2_main_outcome_results.csv")
    df["Dataset"] = df["dataset_name"].map(
        {
            "candidate_gold_strict_v1": "Strict candidate benchmark",
            "candidate_gold_extended_v1": "Extended candidate benchmark",
        }
    )
    df["Model"] = df["model_name"].replace(
        {
            "majority_class": "Majority",
            "rule_baseline": "Rule",
            "tfidf_logreg": "TF-IDF + LogReg",
            "tfidf_linearsvc": "TF-IDF + LinearSVC",
            "tfidf_multinomialnb": "TF-IDF + MNB",
            "current_hybrid_baseline": "Reproducible hybrid",
            "paesc_hybrid": "PAESC hybrid",
        }
    )
    return df[["Dataset", "Model", "accuracy", "macro_f1", "weighted_f1", "macro_f1_ci_low", "macro_f1_ci_high"]].rename(
        columns={
            "accuracy": "Accuracy",
            "macro_f1": "Macro-F1",
            "weighted_f1": "Weighted-F1",
            "macro_f1_ci_low": "CI low",
            "macro_f1_ci_high": "CI high",
        }
    )


def compact_table_3() -> pd.DataFrame:
    df = pd.read_csv(TABLES_DIR / "table3_ablation_leakage_stress.csv")
    keep = df["ablation_setting"].isin(
        ["full_model", "remove_pre_decision_constraint", "remove_evidence_chain", "remove_responsibility_head", "remove_retrieval"]
    )
    df = df[keep].copy()
    df["Dataset"] = df["dataset_name"].map(
        {
            "candidate_gold_strict_v1": "Strict",
            "candidate_gold_extended_v1": "Extended",
        }
    )
    return df[
        ["Dataset", "ablation_setting", "macro_f1", "responsibility_macro_f1", "valid_span_rate", "pre_decision_span_rate", "high_dispute_rate"]
    ].rename(
        columns={
            "ablation_setting": "Setting",
            "macro_f1": "Outcome Macro-F1",
            "responsibility_macro_f1": "Responsibility Macro-F1",
            "valid_span_rate": "Valid span rate",
            "pre_decision_span_rate": "Pre-decision span rate",
            "high_dispute_rate": "High-dispute rate",
        }
    )


def compact_table_4() -> pd.DataFrame:
    df = pd.read_csv(TABLES_DIR / "table4_responsibility_auditability_summary.csv")
    return df[
        [
            "benchmark",
            "responsibility_accuracy",
            "responsibility_macro_f1",
            "candidate_unknown_ratio",
            "strongest_class",
            "strongest_class_f1",
            "weakest_class",
            "weakest_class_f1",
            "valid_span_rate",
            "pre_decision_span_rate",
        ]
    ].rename(
        columns={
            "benchmark": "Benchmark",
            "responsibility_accuracy": "Resp. accuracy",
            "responsibility_macro_f1": "Resp. Macro-F1",
            "candidate_unknown_ratio": "Unknown ratio",
            "strongest_class": "Strongest class",
            "strongest_class_f1": "Strongest F1",
            "weakest_class": "Weakest class",
            "weakest_class_f1": "Weakest F1",
            "valid_span_rate": "Valid span rate",
            "pre_decision_span_rate": "Pre-decision span rate",
        }
    )


def compact_table_5() -> pd.DataFrame:
    df = pd.read_csv(TABLES_DIR / "table5_forensic_claim_boundary.csv")
    return df[
        ["result_group", "run_name", "dataset_name", "model_name", "accuracy", "macro_f1", "audit_status", "claim_tier"]
    ].rename(
        columns={
            "result_group": "Result group",
            "run_name": "Run",
            "dataset_name": "Dataset",
            "model_name": "Model",
            "accuracy": "Accuracy",
            "macro_f1": "Macro-F1",
            "audit_status": "Audit status",
            "claim_tier": "Claim tier",
        }
    )


def build_abstract() -> str:
    return (
        "Construction schedule delay disputes impose substantial transaction costs, documentation burdens, "
        "and governance risks on project organizations. Existing analytics often emphasize outcome prediction "
        "but provide limited support for responsibility reasoning, evidence traceability, or leakage-aware "
        "prospective use. This study proposes DelayDispute Copilot, a management-oriented framework for "
        "construction schedule delay disputes that restricts model inputs to pre-decision information, "
        "reconstructs auditable evidence chains, and generates structured responsibility diagnoses for human "
        "review. The empirical setting uses approximately 4,592 adjudicated cases together with two "
        "machine-assisted candidate benchmarks: a strict benchmark of 250 cases and an extended benchmark "
        "of 500 cases. Under the current audit-ready evaluation pipeline, the PAESC hybrid achieved outcome "
        "macro-F1 scores of 0.552 on the strict candidate benchmark and 0.566 on the extended candidate "
        "benchmark, substantially above traditional TF-IDF baselines. Responsibility diagnosis remained harder, "
        "with known-label macro-F1 scores of 0.180 and 0.196, while evidence-chain auditability remained strong, "
        "with valid-span and pre-decision-span rates above 0.98 on both candidate benchmarks. A forensic audit "
        "further shows that earlier high-scoring historical runs are not suitable as headline claims under the "
        "current audit standard because they lack full manifest and leakage-sentinel support and are partly tied "
        "to materially easier label regimes. The contribution of the study therefore lies not in presenting an "
        "automatic adjudicator, but in establishing a leakage-aware, audit-ready decision-support framework for "
        "dispute triage, documentation diagnosis, and proactive governance in construction delay disputes."
    )


def build_conclusion() -> List[str]:
    return [
        "This paper develops DelayDispute Copilot as a leakage-aware, audit-ready decision-support framework for construction schedule delay disputes. The study shows that auditable evidence-chain reconstruction and structured responsibility outputs can be integrated with outcome prediction under a strict pre-decision information boundary.",
        "The current empirical picture is intentionally cautious. The PAESC hybrid is not the strongest pure predictor in the repository, but it is the main claimable integrated framework because it couples outcome prediction with responsibility diagnosis, evidence tracing, and forensic claim governance. At the same time, responsibility diagnosis remains materially weaker than evidence-chain auditability, which implies that future work should prioritize label governance, coarser responsibility schemas, and expert review protocols before stronger substantive claims are made.",
        "From an engineering management perspective, the practical value of the system lies in helping project teams triage dispute cases earlier, identify documentation and procedural vulnerabilities, and package evidence for human review rather than replacing adjudicators or experts.",
    ]


def add_original_lit_review(doc: Document, base_doc: Document) -> None:
    block = extract_doc_block(base_doc, "II. LITERATURE REVIEW", "III.DATA AND PROBLEM DEFINITION")
    add_paragraph(doc, "II. LITERATURE REVIEW", style="Heading 1", align=WD_ALIGN_PARAGRAPH.LEFT)
    for style_name, text in block[1:]:
        if text.startswith("Table 1."):
            continue
        if style_name.startswith("Heading 2"):
            add_paragraph(doc, text, style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
        else:
            add_paragraph(doc, text)


def add_section_from_md(doc: Document, title: str, section_data: Dict[str, List[str]], level1: str = "Heading 1") -> None:
    add_paragraph(doc, title, style=level1, align=WD_ALIGN_PARAGRAPH.LEFT)
    for sub_title, paragraphs in section_data.items():
        if sub_title != "__main__":
            clean_sub = re.sub(r"^\d+(\.\d+)?\s*", "", sub_title).strip()
            add_paragraph(doc, clean_sub, style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)


def add_references(doc: Document, base_doc: Document) -> None:
    refs: List[str] = []
    capture = False
    for para in base_doc.paragraphs:
        text = para.text.strip()
        if not capture and text.upper().startswith("REFERENCES"):
            capture = True
            continue
        if capture and text:
            refs.append(text)
    add_paragraph(doc, "REFERENCES", style="Heading 1", align=WD_ALIGN_PARAGRAPH.LEFT)
    for ref in refs:
        add_paragraph(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT)


def main() -> None:
    base_doc = Document(str(BASE_DRAFT))
    abstract_doc = Document(str(ABSTRACT_DRAFT))
    md_sections = parse_md_sections(TEXT_DIR / "manuscript_ready_sections.md")

    doc = Document()
    set_document_style(doc)

    title = base_doc.paragraphs[0].text.strip()
    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(15)
    title_run.bold = True

    for para in abstract_doc.paragraphs[1:3]:
        text = para.text.strip()
        if text:
            add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, "Abstract—" + build_abstract(), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(
        doc,
        "Managerial Relevance Statement—" + (TEXT_DIR / "managerial_relevance_statement.md").read_text(encoding="utf-8").strip(),
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Index Terms—Construction disputes; Schedule delay; Claims management; Large language models; Responsibility diagnosis; DelayDispute Copilot.",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    add_section_from_md(doc, "I. INTRODUCTION", md_sections["1. Introduction"])

    add_original_lit_review(doc, base_doc)

    add_paragraph(doc, "III. DATA, PROBLEM DEFINITION, AND EVALUATION DESIGN", style="Heading 1", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "A. Dataset and Candidate Benchmarks", style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in md_sections["2. Method"]["2.1 Research Design and Leakage Control"][:2]:
        add_paragraph(doc, paragraph)
    add_table_from_df(
        doc,
        compact_table_1(),
        "Table 1. Profile of the candidate benchmarks used in the audit-ready evaluation.",
        note="Candidate benchmarks are machine-assisted and should not be described as true human gold labels.",
        font_size=8.2,
    )
    add_figure(
        doc,
        FIGURES_DIR / "fig2_leakage_benchmark_governance.png",
        "Fig. 2. Leakage-control design and candidate-benchmark governance used in the present study.",
        width=6.2,
    )
    add_paragraph(doc, "B. Leakage-Aware Problem Formulation and Evaluation Logic", style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in md_sections["2. Method"]["2.1 Research Design and Leakage Control"][2:]:
        add_paragraph(doc, paragraph)

    add_paragraph(doc, "IV. METHODOLOGY", style="Heading 1", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_figure(
        doc,
        FIGURES_DIR / "fig1_delaydispute_workflow.png",
        "Fig. 1. Leakage-aware workflow of DelayDispute Copilot for construction delay disputes.",
        width=6.3,
    )
    for sub in [
        "2.2 DelayDispute Copilot Framework",
        "2.3 Structured Responsibility Diagnosis",
        "2.4 Evidence-Chain Reconstruction and Auditability",
        "2.5 Forensic Audit and Claim Governance",
    ]:
        add_paragraph(doc, re.sub(r"^\d+(\.\d+)?\s*", "", sub), style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
        for paragraph in md_sections["2. Method"][sub]:
            add_paragraph(doc, paragraph)

    add_paragraph(doc, "V. RESULTS", style="Heading 1", align=WD_ALIGN_PARAGRAPH.LEFT)
    results = md_sections["3. Results"]

    add_paragraph(doc, "A. Dataset Structure and Benchmark Difficulty", style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in results["3.1 Dataset Structure and Benchmark Difficulty"]:
        add_paragraph(doc, paragraph)

    add_paragraph(doc, "B. Main Outcome-Prediction Performance", style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in results["3.2 Main Outcome-Prediction Performance"]:
        add_paragraph(doc, paragraph)
    add_table_from_df(
        doc,
        compact_table_2(),
        "Table 2. Main outcome-prediction results on the strict and extended candidate benchmarks.",
        note="The reproducible hybrid baseline is the strongest predictive comparator, whereas PAESC is retained as the integrated audit-ready framework.",
        font_size=8.0,
    )
    add_figure(
        doc,
        FIGURES_DIR / "fig3_outcome_prediction_comparison.png",
        "Fig. 3. Outcome-prediction performance on the strict and extended candidate benchmarks.",
        width=6.1,
    )
    add_table_from_df(
        doc,
        compact_table_3(),
        "Table 3. Ablation and leakage-stress results for the audit-ready framework.",
        note="Rows removing the pre-decision constraint are stress tests only and are not deployment-valid settings.",
        font_size=7.8,
    )

    add_paragraph(doc, "C. Responsibility Diagnosis and Evidence-Chain Auditability", style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in results["3.3 Responsibility Diagnosis and Evidence-Chain Auditability"]:
        add_paragraph(doc, paragraph)
    add_table_from_df(
        doc,
        compact_table_4(),
        "Table 4. Summary of responsibility-diagnosis performance and evidence-chain auditability.",
        note="Responsibility-diagnosis and evidence-chain outputs are machine-evaluated and audit-ready rather than fully human-validated.",
        font_size=8.1,
    )
    add_figure(
        doc,
        FIGURES_DIR / "fig4_responsibility_auditability.png",
        "Fig. 4. Responsibility-diagnosis performance and evidence-chain auditability on the candidate benchmarks.",
        width=6.1,
    )

    add_paragraph(doc, "D. Leakage Sentinel, Claim Boundary, and Historical Score Recovery", style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
    for key in [
        "3.4 Leakage Sentinel and Inflation Risk",
        "3.5 Why the Legacy High Scores Exist and Why They Are Not the Main Claim",
        "3.6 Root-Cause Analysis of Low Responsibility Performance",
    ]:
        for paragraph in results[key]:
            add_paragraph(doc, paragraph)
    add_table_from_df(
        doc,
        compact_table_5(),
        "Table 5. Forensic claim boundary for historical and current results.",
        note="Tier B results are claimable with caution; Tier C results are historical references only.",
        font_size=8.0,
    )
    add_figure(
        doc,
        FIGURES_DIR / "fig5_forensic_claim_boundary.png",
        "Fig. 5. Forensic audit of leakage-related score inflation risk and result claim boundary.",
        width=6.1,
    )
    add_figure(
        doc,
        FIGURES_DIR / "fig6_representative_case_panel.png",
        "Fig. 6. Representative audit-ready case panel showing pre-decision evidence, predicted outcome, structured responsibility diagnosis, and uncertainty cues.",
        width=6.0,
    )

    add_paragraph(doc, "VI. DISCUSSION", style="Heading 1", align=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in md_sections["4. Discussion"]["__main__"]:
        add_paragraph(doc, paragraph)
    add_paragraph(doc, (TEXT_DIR / "validation_statement.md").read_text(encoding="utf-8").strip())
    add_paragraph(doc, (TEXT_DIR / "limitations_statement.md").read_text(encoding="utf-8").strip())

    add_paragraph(doc, "VII. CONCLUSION", style="Heading 1", align=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in build_conclusion():
        add_paragraph(doc, paragraph)

    add_paragraph(doc, "Supplementary Material Note", style="Heading 2", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(
        doc,
        "Detailed per-class metrics, leakage-sentinel tables, forensic delta audits, expanded ablations, and additional representative cases are preserved in the repository supplementary package under `paper_assets` and should be cited as appendix or online supplementary material rather than treated as main-text evidence.",
    )

    add_references(doc, base_doc)
    doc.save(str(OUTPUT_DOCX))
    print(f"[DONE] merged manuscript written to {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
