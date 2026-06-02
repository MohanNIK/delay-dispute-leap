# -*- coding: utf-8 -*-
"""Build compact Chinese and English Word manuscripts for the MMEC-PAESC 5.5 rerun.

The manuscript generator is intentionally conservative: it uses only verified
prediction-level artifacts from final_eval_55_* and forensic_audit_55_* and keeps
all GPT-5.5/API-unavailable claims inside the audited claim boundary.
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.research_support import load_cfg  # noqa: E402


MODEL_LABELS = {
    "current_hybrid_baseline": "Hybrid baseline",
    "paesc_hybrid": "PAESC",
    "mmec_paesc_55": "MMEC-PAESC proxy",
    "gpt55_direct": "5.5 direct proxy",
}
DATASET_LABELS = {
    "candidate_gold_strict_v1": "Strict candidate",
    "candidate_gold_extended_v1": "Extended candidate",
}


def latest_prefixed_dir(root: Path, prefix: str) -> Path:
    runs = sorted([p for p in root.glob(f"{prefix}*") if p.is_dir()])
    if not runs:
        raise RuntimeError(f"No run found for prefix {prefix}")
    return runs[-1]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_style(doc: Document, cn: bool = False) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun" if cn else "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.2)


def add_title(doc: Document, text: str, cn: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "SimSun" if cn else "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r.font.size = Pt(14.5)


def add_heading(doc: Document, text: str, level: int = 1, cn: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "SimSun" if cn else "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r.font.size = Pt(12 if level == 1 else 10.5)


def add_para(doc: Document, text: str, cn: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cn else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.72)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "SimSun" if cn else "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r.font.size = Pt(10.2)


def add_caption(doc: Document, text: str, cn: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "SimSun" if cn else "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r.font.size = Pt(8.8)


def add_figure(doc: Document, path: Path, caption: str, cn: bool = False) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(5.75))
    add_caption(doc, caption, cn=cn)


def add_df_table(doc: Document, df: pd.DataFrame, caption: str, cn: bool = False, max_rows: int = 12) -> None:
    add_caption(doc, caption, cn=cn)
    df = df.head(max_rows).copy()
    display = df.copy()
    for col in display.columns:
        display[col] = display[col].map(lambda v: f"{v:.3f}" if isinstance(v, float) else str(v))

    # Render tables as embedded PNGs for stable cross-renderer layout. The
    # editable source tables remain exported separately under paper_assets.
    img_dir = PROJECT_ROOT / "paper_assets" / "manuscript" / "table_images"
    img_dir.mkdir(parents=True, exist_ok=True)
    idx = getattr(add_df_table, "_counter", 0) + 1
    setattr(add_df_table, "_counter", idx)
    img_path = img_dir / f"table_{idx}_{'cn' if cn else 'en'}.png"
    fig_w = 6.8
    fig_h = max(1.0, 0.32 * (len(display) + 1))
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis("off")
    table = ax.table(
        cellText=display.values,
        colLabels=display.columns,
        cellLoc="center",
        loc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(7.2)
    table.scale(1.0, 1.18)
    for (row, col), cell in table.get_celld().items():
        cell.set_linewidth(0.35)
        if row == 0:
            cell.set_facecolor("#EDEDED")
            cell.get_text().set_weight("bold")
        if col in (0, 1) and row > 0:
            cell.get_text().set_ha("left")
    fig.savefig(img_path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Inches(6.15))
    doc.add_paragraph()


def load_results(run_dir: Path, audit_dir: Path) -> Dict[str, object]:
    baseline = pd.read_csv(run_dir / "baseline_comparison.csv", encoding="utf-8-sig")
    claims = pd.read_csv(audit_dir / "claim_tiering.csv", encoding="utf-8-sig")
    chain = pd.read_csv(run_dir / "evidence_chain_eval.csv", encoding="utf-8-sig")
    resp = pd.read_csv(run_dir / "responsibility_eval.csv", encoding="utf-8-sig")
    manifest = json.loads((run_dir / "run_manifest.json").read_text(encoding="utf-8"))

    def macro(dataset: str, model: str) -> float:
        sub = baseline[(baseline["dataset_name"] == dataset) & (baseline["model_name"] == model)]
        return float(sub["macro_f1"].iloc[0]) if not sub.empty else 0.0

    def acc(dataset: str, model: str) -> float:
        sub = baseline[(baseline["dataset_name"] == dataset) & (baseline["model_name"] == model)]
        return float(sub["accuracy"].iloc[0]) if not sub.empty else 0.0

    chain_mmec = chain[chain.get("model_name", "") == "mmec_paesc_55"] if "model_name" in chain.columns else chain
    if chain_mmec.empty:
        chain_mmec = chain
    return {
        "baseline": baseline,
        "claims": claims,
        "chain": chain,
        "resp": resp,
        "manifest": manifest,
        "strict_mmec": macro("candidate_gold_strict_v1", "mmec_paesc_55"),
        "extended_mmec": macro("candidate_gold_extended_v1", "mmec_paesc_55"),
        "strict_mmec_acc": acc("candidate_gold_strict_v1", "mmec_paesc_55"),
        "extended_mmec_acc": acc("candidate_gold_extended_v1", "mmec_paesc_55"),
        "strict_paesc": macro("candidate_gold_strict_v1", "paesc_hybrid"),
        "extended_paesc": macro("candidate_gold_extended_v1", "paesc_hybrid"),
        "strict_hybrid": macro("candidate_gold_strict_v1", "current_hybrid_baseline"),
        "extended_hybrid": macro("candidate_gold_extended_v1", "current_hybrid_baseline"),
        "strict_hybrid_acc": acc("candidate_gold_strict_v1", "current_hybrid_baseline"),
        "extended_hybrid_acc": acc("candidate_gold_extended_v1", "current_hybrid_baseline"),
        "api_counts": manifest.get("api_status_counts", {}),
        "valid_span": float(chain_mmec["valid_span_rate"].mean()),
        "pre_span": float(chain_mmec["pre_decision_span_rate"].mean()),
    }


def selected_tables(results: Dict[str, object]) -> List[Tuple[str, pd.DataFrame]]:
    baseline = results["baseline"].copy()
    models = ["current_hybrid_baseline", "paesc_hybrid", "mmec_paesc_55", "gpt55_direct"]
    table2 = baseline[baseline["model_name"].isin(models)].copy()
    table2["Benchmark"] = table2["dataset_name"].map(DATASET_LABELS)
    table2["Model"] = table2["model_name"].map(MODEL_LABELS)
    table2 = table2[["Benchmark", "Model", "accuracy", "macro_f1", "weighted_f1"]]
    table2 = table2.rename(columns={"accuracy": "Acc.", "macro_f1": "Macro-F1", "weighted_f1": "Weighted-F1"})

    claims = results["claims"].copy()
    claims = claims[claims["model_name"].isin(["paesc_hybrid", "current_hybrid_baseline", "gpt55_direct", "mmec_paesc_55"])].copy()
    claims["Model"] = claims["model_name"].map(MODEL_LABELS).fillna(claims["model_name"])
    claims["Tier"] = claims["claim_tier"].str.replace("Tier B: claimable with caution", "Tier B", regex=False)
    claims["Tier"] = claims["Tier"].str.replace("Tier C: historical/reference only", "Tier C", regex=False)
    table_claim = claims[["Model", "accuracy_recomputed", "macro_f1_recomputed", "Tier"]]
    table_claim = table_claim.rename(columns={"accuracy_recomputed": "Acc.", "macro_f1_recomputed": "Macro-F1"})
    return [("Main outcome results", table2), ("Forensic claim boundary", table_claim)]


def build_en(run_dir: Path, audit_dir: Path, out_path: Path) -> None:
    res = load_results(run_dir, audit_dir)
    doc = Document()
    set_doc_style(doc)
    add_title(doc, "DelayDispute Copilot 5.5: Management-Mechanism Evidence Calibration for Construction Schedule Delay Disputes")
    add_para(doc, f"Abstract. This revised manuscript reports a controlled audit-first rerun of DelayDispute Copilot using the MMEC-PAESC branch. The rerun preserves the pre-decision input boundary, the existing candidate benchmarks, and the original label definitions. MMEC-PAESC achieved macro-F1 scores of {res['strict_mmec']:.4f} on the strict candidate benchmark and {res['extended_mmec']:.4f} on the extended candidate benchmark. The strongest reproducible pure predictor remains the hybrid baseline with macro-F1 scores of {res['strict_hybrid']:.4f} and {res['extended_hybrid']:.4f}. Evidence-chain auditability remains strong, with valid-span and pre-decision-span rates of {res['valid_span']:.3f} and {res['pre_span']:.3f}. The API status record is {res['api_counts']}; therefore, the 5.5 rows are treated as local MMEC mechanism proxies unless api_available rows are later produced.")
    add_heading(doc, "I. Introduction")
    add_para(doc, "Construction schedule delay disputes are governance problems involving entitlement, causality, procedural compliance, documentation integrity, and responsibility allocation. This manuscript positions DelayDispute Copilot as an audit-ready decision-support framework for project teams rather than an autonomous adjudicator.")
    add_heading(doc, "II. Method: MMEC-PAESC")
    add_para(doc, "MMEC-PAESC adds Management Mechanism Evidence Calibration to the existing PAESC pipeline. The mechanism layer extracts documentation-gap, procedural-risk, causality-ambiguity, concurrency-risk, critical-path-support, and negotiation-readiness indices from pre-decision text only. These indices calibrate outcome boundaries and provide managerial diagnosis without changing the candidate benchmark or label definitions.")
    add_figure(doc, PROJECT_ROOT / "paper_assets" / "figures" / "fig1_delaydispute_workflow.png", "Fig. 1. Leakage-aware DelayDispute Copilot workflow.")
    for caption, df in selected_tables(res):
        add_df_table(doc, df, f"Table. {caption}.")
    add_heading(doc, "III. Results")
    add_para(doc, f"The 5.5-controlled rerun did not improve headline predictive accuracy. Relative to the earlier audit-ready PAESC values of {res['strict_paesc']:.4f}/{res['extended_paesc']:.4f}, the MMEC-PAESC proxy reports {res['strict_mmec']:.4f}/{res['extended_mmec']:.4f} on the strict/extended candidate benchmarks. The strongest reproducible pure predictor remains the hybrid baseline at {res['strict_hybrid']:.4f}/{res['extended_hybrid']:.4f}. Thus, the safe interpretation is mechanism-oriented: MMEC strengthens the managerial mechanism representation, but it is not a verified GPT-5.5 accuracy gain.")
    add_heading(doc, "IV. Discussion and Managerial Relevance")
    add_para(doc, "The main value of the revision is management mechanism reconstruction. The system identifies whether a dispute is constrained by documentation gaps, procedural noncompliance, ambiguous causality, concurrent delay narratives, or insufficient critical-path support. These outputs support dispute triage, record supplementation, negotiation preparation, and human expert review.")
    add_heading(doc, "V. Claim Boundary")
    add_para(doc, "Candidate benchmarks are machine-assisted and are not described as human gold labels. Responsibility diagnosis and evidence-chain outputs are audit-ready and machine-evaluated rather than fully human-validated. Legacy high scores and API-unavailable 5.5 rows remain reference results unless they pass traceability, metric recomputation, and leakage-sentinel checks.")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def build_cn(run_dir: Path, audit_dir: Path, out_path: Path) -> None:
    res = load_results(run_dir, audit_dir)
    doc = Document()
    set_doc_style(doc, cn=True)
    add_title(doc, "DelayDispute Copilot 5.5：面向建筑工期延误纠纷的管理机制证据校准框架", cn=True)
    add_para(doc, f"摘要：本文在不改变测试集、不改变标签定义、不引入裁决后信息泄漏的前提下，对 DelayDispute Copilot 进行了审计优先的 5.5 分支重跑。MMEC-PAESC 在严格候选基准和扩展候选基准上的 macro-F1 分别为 {res['strict_mmec']:.4f} 和 {res['extended_mmec']:.4f}。当前最强可复现纯预测基线仍为 hybrid baseline，其 macro-F1 分别为 {res['strict_hybrid']:.4f} 和 {res['extended_hybrid']:.4f}。证据链 valid-span rate 与 pre-decision-span rate 分别为 {res['valid_span']:.3f} 和 {res['pre_span']:.3f}。API 状态记录为 {res['api_counts']}；因此，若不存在 api_available 行，本轮 5.5 结果只能表述为本地 MMEC 机制代理结果，不能写成真实 GPT-5.5 精度提升。", cn=True)
    add_heading(doc, "一、研究定位", cn=True)
    add_para(doc, "工期延误纠纷不是单纯的文本分类问题，而是由合同权利基础、因果关系、程序履约、证据完整性和责任配置共同构成的工程管理治理问题。本文将系统定位为面向项目前端治理和争议评估的人机协同决策支持工具，而不是自动裁判系统。", cn=True)
    add_heading(doc, "二、方法：MMEC-PAESC", cn=True)
    add_para(doc, "MMEC 即 Management Mechanism Evidence Calibration，在 PAESC 的基础上增加管理机制校准层。系统仅从裁决前文本中抽取证据缺口、程序履约风险、因果模糊度、并发延误风险、关键线路支持和谈判准备度等指标，并将这些指标用于校准结果边界和输出管理诊断。该分支不改变候选基准、不改变标签定义，也不使用裁决后文本作为推理输入。", cn=True)
    add_figure(doc, PROJECT_ROOT / "paper_assets" / "figures" / "fig1_delaydispute_workflow.png", "图 1. DelayDispute Copilot 的泄漏控制与证据链工作流。", cn=True)
    for caption, df in selected_tables(res):
        cn_caption = "主结果表" if "Main" in caption else "法证审计 claim boundary"
        add_df_table(doc, df, f"表. {cn_caption}。", cn=True)
    add_heading(doc, "三、结果", cn=True)
    add_para(doc, f"本轮 5.5 受控重跑没有提升 headline 预测精度。相对于原 PAESC 审计主方法的 {res['strict_paesc']:.4f}/{res['extended_paesc']:.4f}，MMEC-PAESC 代理分支在 strict/extended 候选基准上的结果为 {res['strict_mmec']:.4f}/{res['extended_mmec']:.4f}。当前最强可复现纯预测基线仍为 hybrid baseline，其结果为 {res['strict_hybrid']:.4f}/{res['extended_hybrid']:.4f}。因此，安全写法不是“5.5 提高了精度”，而是“管理机制重构增强了审计与解释维度，但未形成经验证的 GPT-5.5 精度提升”。", cn=True)
    add_heading(doc, "四、管理学贡献", cn=True)
    add_para(doc, "本轮修订的价值在于管理机制重构：系统能够识别争议中的证据缺口、程序脆弱性、因果模糊、并发延误叙事和关键线路支持不足，并将其转化为项目团队可以理解的治理信号。这些输出可服务于争议分流、资料补强、谈判准备、责任复盘和专家复核。", cn=True)
    add_heading(doc, "五、结论与边界", cn=True)
    add_para(doc, "候选基准仍是机器辅助候选标签，不是人工金标；责任诊断和证据链是可审计、可复核的机器输出，不是完全人工验证结论。历史高分和 API 不可用的 5.5 行只能作为参考结果，不能进入主文 headline claim。", cn=True)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", default="config/research_v2_55.yaml")
    ap.add_argument("--run-dir", default="")
    ap.add_argument("--audit-dir", default="")
    args = ap.parse_args()
    cfg = load_cfg(PROJECT_ROOT / args.config)
    run_dir = Path(args.run_dir) if args.run_dir else latest_prefixed_dir(PROJECT_ROOT / cfg["paths"]["final_eval_root"], "final_eval_55_")
    audit_dir = Path(args.audit_dir) if args.audit_dir else latest_prefixed_dir(PROJECT_ROOT / cfg["paths"]["final_eval_root"], "forensic_audit_55_")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_root = PROJECT_ROOT / "paper_assets" / "manuscript"
    build_en(run_dir, audit_dir, out_root / f"DelayDispute_Copilot_55_EN_{stamp}.docx")
    build_cn(run_dir, audit_dir, out_root / f"DelayDispute_Copilot_55_CN_{stamp}.docx")
    print(f"[DONE] bilingual manuscripts in {out_root}")


if __name__ == "__main__":
    main()
