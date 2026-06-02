# -*- coding: utf-8 -*-
"""Leakage-aware data structuring pipeline for DelayDispute Copilot.

This script keeps the repository aligned with the paper framing:
- construction schedule delay disputes
- pre-decision information for model inputs
- post-decision information only for label derivation / evaluation support
- auditable structured case records

Stages
------
parse:
    DOCX -> parsed JSON in data/2_parsed_json

enrich:
    parsed JSON -> structured leakage-aware case JSON in data/3_structured_cases

prepare_labels:
    refreshes meta labels with domain flags and structured-case availability

all:
    parse (if needed) + enrich + prepare_labels
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
from docx import Document
from tqdm import tqdm

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.research_support import (
    DEFAULT_CFG,
    build_structured_case,
    detect_domain_case,
    ensure_dir,
    json_dump,
    load_cfg,
    normalize_label,
    read_csv_flexible,
)

PATTERNS = {
    "facts": [r"事实与理由", r"本院经审理认定", r"本院查明", r"案件事实", r"事实认定"],
    "issues": [r"争议焦点", r"争议在于", r"本案争议焦点", r"主要争议"],
    "reasoning": [r"本院认为", r"法院认为", r"裁判理由", r"本庭认为"],
    "decision": [r"判决如下", r"裁定如下", r"判令如下", r"判决主文"],
}


def file_sha1(path: Path) -> str:
    h = hashlib.sha1()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def normalize_text(text: str) -> str:
    text = (text or "").replace("\u3000", " ").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_docx_text(docx_path: Path, include_tables: bool = True) -> str:
    doc = Document(str(docx_path))
    chunks: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            chunks.append(t)
    if include_tables:
        for tb in doc.tables:
            for row in tb.rows:
                row_texts = []
                for cell in row.cells:
                    ct = re.sub(r"\s+", " ", (cell.text or "").strip())
                    if ct:
                        row_texts.append(ct)
                if row_texts:
                    chunks.append(" | ".join(row_texts))
    return normalize_text("\n".join(chunks))


def _find_positions(text: str) -> Dict[str, int]:
    positions: Dict[str, int] = {}
    for key, pattern_list in PATTERNS.items():
        for pattern in pattern_list:
            m = re.search(pattern, text)
            if m:
                positions[key] = m.start()
                break
    return positions


def split_sections(text: str) -> Dict[str, str]:
    text = normalize_text(text)
    positions = _find_positions(text)
    if len(positions) < 2:
        return {
            "facts": "",
            "issues": "",
            "reasoning": "",
            "decision": "",
            "full_text": text,
            "segmentation_note": "fallback_full_text_only",
        }

    ordered = sorted(positions.items(), key=lambda x: x[1])
    sections = {"facts": "", "issues": "", "reasoning": "", "decision": "", "full_text": text, "segmentation_note": "heading_based_v3"}
    starts = [p for _, p in ordered]
    for i, (sec, start) in enumerate(ordered):
        end = starts[i + 1] if i + 1 < len(starts) else len(text)
        sections[sec] = text[start:end].strip()
    return sections


def step1_parse_docx(cfg: Dict, incremental: bool = True, overwrite: bool = False) -> Path:
    raw_dir = PROJECT_ROOT / cfg["paths"].get("raw_docx_dir", "data/0_raw_docx")
    out_dir = ensure_dir(PROJECT_ROOT / cfg["paths"]["parsed_json_dir"])
    txt_dir = ensure_dir(PROJECT_ROOT / cfg["paths"].get("raw_text_dir", "data/1_raw_text"))
    if not raw_dir.exists():
        print(f"[INFO] raw docx dir not found: {raw_dir}. Skip parse stage.")
        return out_dir

    files = sorted(raw_dir.glob("*.docx"))
    for fp in tqdm(files, desc="parse_docx"):
        case_id = file_sha1(fp)[:12]
        out_json = out_dir / f"{case_id}.json"
        out_txt = txt_dir / f"{case_id}.txt"
        if incremental and out_json.exists() and not overwrite:
            continue
        text = read_docx_text(fp)
        parsed = {
            "case_id": case_id,
            "source_file": fp.name,
            "sha1": file_sha1(fp),
            "sections": split_sections(text),
        }
        json_dump(out_json, parsed)
        out_txt.write_text(text, encoding="utf-8")
    return out_dir


def enrich_structured_cases(cfg: Dict, overwrite: bool = False) -> Tuple[Path, pd.DataFrame]:
    parsed_dir = PROJECT_ROOT / cfg["paths"]["parsed_json_dir"]
    out_dir = ensure_dir(PROJECT_ROOT / cfg["paths"]["structured_case_dir"])
    rows = []
    for fp in tqdm(sorted(parsed_dir.glob("*.json")), desc="build_structured_cases"):
        structured_fp = out_dir / fp.name
        if structured_fp.exists() and not overwrite:
            try:
                obj = json.loads(structured_fp.read_text(encoding="utf-8"))
                rows.append({
                    "case_id": obj.get("case_id"),
                    "source_file": obj.get("source_file", ""),
                    "case_year": obj.get("case_year"),
                    "is_domain_case": int(detect_domain_case(obj.get("source_file", ""), obj.get("pre_decision_text", ""))),
                    "potential_leakage_flag": obj.get("potential_leakage_flag", 0),
                    "pre_post_split_confidence": obj.get("pre_post_split_confidence", 0.0),
                })
                continue
            except Exception:
                pass
        parsed = json.loads(fp.read_text(encoding="utf-8"))
        structured = build_structured_case(parsed)
        structured["is_domain_case"] = int(detect_domain_case(structured.get("source_file", ""), structured.get("pre_decision_text", "")))
        json_dump(structured_fp, structured)
        rows.append({
            "case_id": structured.get("case_id"),
            "source_file": structured.get("source_file", ""),
            "case_year": structured.get("case_year"),
            "is_domain_case": structured.get("is_domain_case", 0),
            "potential_leakage_flag": structured.get("potential_leakage_flag", 0),
            "pre_post_split_confidence": structured.get("pre_post_split_confidence", 0.0),
        })

    index_df = pd.DataFrame(rows).drop_duplicates(subset=["case_id"])
    ensure_dir((PROJECT_ROOT / "data" / "meta"))
    index_df.to_csv(PROJECT_ROOT / "data" / "meta" / "structured_case_index.csv", index=False, encoding="utf-8-sig")
    return out_dir, index_df


def prepare_labels(cfg: Dict) -> Path:
    labels_csv = PROJECT_ROOT / cfg["paths"]["meta_labels_csv"]
    structured_index = read_csv_flexible(PROJECT_ROOT / "data" / "meta" / "structured_case_index.csv")
    structured_index["case_id"] = structured_index.get("case_id", pd.Series(dtype=str)).astype(str)

    if labels_csv.exists():
        labels = read_csv_flexible(labels_csv)
        labels["case_id"] = labels["case_id"].astype(str)
    else:
        rows = []
        structured_dir = PROJECT_ROOT / cfg["paths"]["structured_case_dir"]
        for fp in tqdm(sorted(structured_dir.glob("*.json")), desc="bootstrap_labels"):
            obj = json.loads(fp.read_text(encoding="utf-8"))
            post = obj.get("post_decision_text", "")
            label = "unknown"
            if "部分支持" in post or ("其余" in post and "驳回" in post):
                label = "partial"
            elif "不予支持" in post or "驳回" in post:
                label = "not_support"
            elif "支持" in post or "判令" in post:
                label = "support"
            rows.append({
                "case_id": obj.get("case_id"),
                "source_file": obj.get("source_file", ""),
                "eot_source_note": "post_bootstrap",
                "eot_label": label,
                "eot_evidence": post[:240],
                "cost_label": "unknown",
                "cost_evidence": "",
            })
        labels = pd.DataFrame(rows)

    labels["eot_label"] = labels["eot_label"].map(normalize_label)
    labels = labels.merge(structured_index, on=["case_id", "source_file"], how="left")
    labels["is_domain_case"] = labels["is_domain_case"].fillna(0).astype(int)
    labels["pre_post_split_confidence"] = labels["pre_post_split_confidence"].fillna(0.0)
    labels.to_csv(labels_csv, index=False, encoding="utf-8-sig")

    domain_labels = labels[labels["is_domain_case"] == 1].copy()
    domain_path = PROJECT_ROOT / "data" / "meta" / "labels_step2_domain.csv"
    domain_labels.to_csv(domain_path, index=False, encoding="utf-8-sig")
    return labels_csv


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", type=str, default="config/research_v1.yaml")
    ap.add_argument("--stage", choices=["parse", "enrich", "prepare_labels", "all"], default="all")
    ap.add_argument("--overwrite", action="store_true")
    ap.add_argument("--no_incremental", action="store_true")
    args = ap.parse_args()

    cfg = load_cfg(PROJECT_ROOT / args.config)
    if args.stage in ("parse", "all"):
        step1_parse_docx(cfg, incremental=not args.no_incremental, overwrite=args.overwrite)
    if args.stage in ("enrich", "all"):
        enrich_structured_cases(cfg, overwrite=args.overwrite)
    if args.stage in ("prepare_labels", "all"):
        prepare_labels(cfg)

    print("[PIPELINE DONE]")


if __name__ == "__main__":
    main()
