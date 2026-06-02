# -*- coding: utf-8 -*-
"""Append true LLM Copilot results and IEEE-TEM notes to the paper4 Word draft."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
RUN_DIR = PROJECT_ROOT / "results" / "true_llm_copilot_20260519_112238"
FIG_DIR = PROJECT_ROOT / "paper_assets" / "figures_true_llm"
OUT_DIR = PROJECT_ROOT / "paper_assets" / "manuscript"
TEXT_DIR = PROJECT_ROOT / "paper_assets" / "text"


def find_source_docx() -> Path:
    candidates = list(Path(r"C:\Users\pig\Desktop").glob("paper4*.docx"))
    if not candidates:
        raise FileNotFoundError("Cannot find paper4*.docx on Desktop")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def set_run_font(run, name: str = "Times New Roman", east_asia: str = "宋体", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Key point: ")
    set_run_font(run, bold=True)
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(3)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", size=14 if level == 1 else 12, bold=True)


def add_table(doc: Document, df: pd.DataFrame, columns: List[str], title: str) -> None:
    add_para(doc, title)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = col
        set_cell_shading(hdr[i], "D9EAF7")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, east_asia="黑体", size=9, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                val = f"{val:.3f}"
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=8)


def add_figure(doc: Document, path: Path, caption: str, analysis: str) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=8, bold=True)
    add_para(doc, analysis)


def load_key_metrics() -> Dict[str, float]:
    metrics = json.loads((RUN_DIR / "metrics_main.json").read_text(encoding="utf-8"))
    ext = metrics["candidate_gold_evaluation"]["candidate_gold_extended_v1"]["true_qwen_direct"]
    strict = metrics["candidate_gold_evaluation"]["candidate_gold_strict_v1"]["true_qwen_direct"]
    resp = pd.read_csv(RUN_DIR / "responsibility_summary.csv", encoding="utf-8-sig")
    ev = pd.read_csv(RUN_DIR / "evidence_chain_eval.csv", encoding="utf-8-sig")
    comp = pd.read_csv(RUN_DIR / "model_comparison.csv", encoding="utf-8-sig")
    mech = pd.read_csv(RUN_DIR / "managerial_mechanisms.csv", encoding="utf-8-sig")
    out = {
        "strict_macro_f1": strict["macro_f1"],
        "strict_accuracy": strict["accuracy"],
        "extended_macro_f1": ext["macro_f1"],
        "extended_accuracy": ext["accuracy"],
        "api_success_rate": ext["api_success_rate"],
        "strict_resp_folded_f1": float(resp[resp["dataset_name"].eq("candidate_gold_strict_v1")]["folded_macro_f1"].iloc[0]),
        "extended_resp_folded_f1": float(resp[resp["dataset_name"].eq("candidate_gold_extended_v1")]["folded_macro_f1"].iloc[0]),
        "extended_resp_fine_f1": float(resp[resp["dataset_name"].eq("candidate_gold_extended_v1")]["fine_macro_f1"].iloc[0]),
        "valid_span_rate": float(ev["valid_span_rate"].mean()),
        "pre_decision_span_rate": float(ev["pre_decision_span_rate"].mean()),
        "role_coverage_rate": float(ev["role_coverage_rate"].mean()),
        "hybrid_extended_macro_f1": float(comp[(comp["dataset_name"].eq("candidate_gold_extended_v1")) & (comp["model_name"].eq("current_hybrid_baseline"))]["macro_f1"].iloc[0]),
        "paesc_extended_macro_f1": float(comp[(comp["dataset_name"].eq("candidate_gold_extended_v1")) & (comp["model_name"].eq("paesc_hybrid"))]["macro_f1"].iloc[0]),
        "procedure_risk": float(mech[mech["dataset_name"].eq("candidate_gold_extended_v1")]["procedural_compliance_risk"].mean()),
        "negotiation_readiness": float(mech[mech["dataset_name"].eq("candidate_gold_extended_v1")]["negotiation_readiness_score"].mean()),
    }
    return out


def write_markdown_summary(metrics: Dict[str, float]) -> Path:
    out = TEXT_DIR / "true_llm_copilot_result_analysis_ieeetem.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        f"""# True LLM Copilot Result Analysis and IEEE-TEM Claim Boundary

## Verified API-backed result
- API branch: DashScope/Qwen `qwen-plus`.
- API success rate on the candidate benchmark: {metrics['api_success_rate']:.3f}.
- Extended outcome macro-F1: {metrics['extended_macro_f1']:.4f}.
- Strict outcome macro-F1: {metrics['strict_macro_f1']:.4f}.
- Extended fine-grained responsibility macro-F1: {metrics['extended_resp_fine_f1']:.4f}.
- Extended folded responsibility macro-F1: {metrics['extended_resp_folded_f1']:.4f}.
- Mean valid-span rate: {metrics['valid_span_rate']:.4f}.
- Mean pre-decision-span rate: {metrics['pre_decision_span_rate']:.4f}.
- Mean role-coverage rate: {metrics['role_coverage_rate']:.4f}.

## Academic-paper reasonableness check
The direct LLM branch is not suitable as the headline predictor because its outcome macro-F1 is lower than both PAESC and the current hybrid baseline. Its manuscript value is different: it provides true API-backed evidence-chain reconstruction, structured responsibility diagnosis, and management mechanism indicators under a pre-decision-only constraint. The safest academic framing is therefore to keep PAESC as the audit-ready integrated framework, retain the current hybrid baseline as the strongest pure predictor, and present true Qwen direct as an exploratory Copilot layer that strengthens traceability and management interpretation rather than raw prediction accuracy.

## IEEE-TEM minimal style adaptation
The paper should not be written as an LLM adjudication paper. It should be written as an engineering management decision-support paper. The recommended terminology is: leakage-aware decision support, human-in-the-loop review, candidate benchmark, audit-ready evidence chain, responsibility signal, and management mechanism indicator. Avoid: human gold, automatic adjudicator, fully validated responsibility diagnosis, and LLM outperforms all baselines.
""",
        encoding="utf-8",
    )
    return out


def build_docx() -> Path:
    src = find_source_docx()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"paper4_IEEE_TEM_true_LLM_minimal_revision_{timestamp}.docx"
    doc = Document(src)
    style_doc(doc)
    metrics = load_key_metrics()
    write_markdown_summary(metrics)

    doc.add_page_break()
    add_heading(doc, "Revision Addendum: True LLM Copilot Results and IEEE-TEM Claim Boundary", 1)
    add_para(
        doc,
        "This addendum minimally updates the manuscript with the newly verified true LLM Copilot branch. The original manuscript is preserved; the new material clarifies what the API-backed LLM can and cannot support. The candidate benchmarks remain machine-assisted candidate benchmarks rather than human gold labels.",
    )

    add_heading(doc, "1. Verified True LLM Copilot Results", 2)
    add_bullet(doc, f"True Qwen direct achieved 100% API completion on the 500 unique candidate cases; no rule fallback was used for headline metrics.")
    add_bullet(doc, f"Outcome prediction remained weaker than the existing baselines: extended macro-F1 = {metrics['extended_macro_f1']:.4f}, compared with PAESC = {metrics['paesc_extended_macro_f1']:.4f} and current hybrid baseline = {metrics['hybrid_extended_macro_f1']:.4f}.")
    add_bullet(doc, f"Evidence-chain auditability was substantially stronger than direct outcome classification: mean valid-span rate = {metrics['valid_span_rate']:.4f}, pre-decision-span rate = {metrics['pre_decision_span_rate']:.4f}, and role coverage = {metrics['role_coverage_rate']:.4f}.")
    add_bullet(doc, f"Responsibility diagnosis remains a difficult, audit-ready task: extended fine-grained macro-F1 = {metrics['extended_resp_fine_f1']:.4f}; folded-schema macro-F1 = {metrics['extended_resp_folded_f1']:.4f}.")

    comp = pd.read_csv(RUN_DIR / "model_comparison.csv", encoding="utf-8-sig")
    comp = comp[comp["dataset_name"].eq("candidate_gold_extended_v1")][["model_name", "accuracy", "macro_f1", "weighted_f1", "claim_note"]]
    add_table(doc, comp, ["model_name", "accuracy", "macro_f1", "weighted_f1", "claim_note"], "Table A1. Extended candidate benchmark comparison from prediction-level artifacts.")

    add_heading(doc, "2. Figure-Based Interpretation", 2)
    add_figure(
        doc,
        FIG_DIR / "fig_true1_model_comparison_macro_f1.png",
        "Fig. A1. Outcome prediction comparison between the verified true LLM branch and reproducible baselines.",
        "Interpretation: the direct LLM branch should not replace PAESC or the current hybrid baseline for headline prediction. Its extended macro-F1 is substantially lower, so the manuscript should frame it as a true API-backed Copilot diagnostic layer rather than as a more accurate classifier.",
    )
    add_figure(
        doc,
        FIG_DIR / "fig_true2_qwen_confusion_matrices.png",
        "Fig. A2. Confusion matrices for true Qwen direct on the strict and extended candidate benchmarks.",
        "Interpretation: the direct LLM output tends to over-select partial support and has limited discrimination for not-support cases. This supports a human-in-the-loop design in which LLM outputs are reviewed as structured signals rather than accepted as final adjudicative labels.",
    )
    add_figure(
        doc,
        FIG_DIR / "fig_true3_evidence_auditability.png",
        "Fig. A3. Evidence-chain auditability of the true LLM Copilot branch.",
        "Interpretation: the strongest verified contribution of the true LLM branch is evidence-chain reconstruction. Most cited spans are valid and located inside the pre-decision text, which supports auditability and managerial review even when label accuracy is limited.",
    )
    add_figure(
        doc,
        FIG_DIR / "fig_true4_responsibility_management.png",
        "Fig. A4. Responsibility diagnosis and management mechanism indicators.",
        "Interpretation: responsibility classification remains weak under fine-grained labels, but the folded schema and management indicators are more useful for project governance. This is aligned with IEEE-TEM positioning because the contribution is decision support, evidence governance, and risk triage rather than autonomous legal prediction.",
    )

    add_heading(doc, "3. Academic-Paper Reasonableness Check", 2)
    add_para(
        doc,
        "The paper is reasonable if it is positioned as an engineering management decision-support study rather than a pure NLP benchmark. The empirical evidence does not support a claim that direct LLM prediction is superior. It does support a narrower and stronger claim: under a pre-decision-only boundary, DelayDispute Copilot can convert unstructured dispute materials into auditable evidence chains, structured responsibility signals, and management mechanism indicators for expert review.",
    )
    add_bullet(doc, "Safe claim: PAESC remains the main audit-ready integrated framework; the current hybrid baseline remains the strongest reproducible pure predictor.")
    add_bullet(doc, "Safe claim: the true Qwen branch verifies that an API-backed LLM can generate traceable evidence-chain and management mechanism outputs at scale.")
    add_bullet(doc, "Caution: candidate benchmarks are machine-assisted and should not be described as human gold.")
    add_bullet(doc, "Caution: responsibility diagnosis is audit-ready and machine-evaluated, not fully human-validated.")
    add_bullet(doc, "Caution: direct LLM outcome prediction is exploratory and should not be a headline performance claim.")

    add_heading(doc, "4. IEEE-TEM Minimal Style Adaptation", 2)
    add_para(
        doc,
        "Recommended replacement for the main result narration:",
        style=None,
    )
    add_para(
        doc,
        f"Using only pre-decision information, the audit-ready PAESC framework achieved an extended candidate-benchmark macro-F1 of {metrics['paesc_extended_macro_f1']:.4f}, while the strongest reproducible pure predictor, the current hybrid baseline, achieved {metrics['hybrid_extended_macro_f1']:.4f}. The newly verified true Qwen direct branch achieved an extended macro-F1 of {metrics['extended_macro_f1']:.4f}, indicating that direct LLM prompting alone is not sufficient as a headline classifier under the present candidate-benchmark setting. However, the same API-backed branch produced substantially stronger audit artifacts: cited evidence spans were valid and pre-decision-constrained at a mean rate of {metrics['valid_span_rate']:.4f}, with mean role coverage of {metrics['role_coverage_rate']:.4f}. These findings support the paper's central management claim: the value of DelayDispute Copilot lies in leakage-aware, human-in-the-loop dispute governance rather than autonomous adjudication.",
    )
    add_para(
        doc,
        "Recommended terminology edits: replace 'gold set' with 'machine-assisted candidate benchmark'; replace 'automatic responsibility determination' with 'audit-ready responsibility signal'; replace 'LLM predicts the judgment' with 'LLM-supported evidence-chain decision support'; replace 'model replaces expert judgment' with 'human-in-the-loop project governance support'.",
    )

    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_docx())
