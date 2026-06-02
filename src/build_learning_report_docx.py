# -*- coding: utf-8 -*-
"""Build an English paper-style learning report as DOCX and Markdown."""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Optional

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parent.parent


def latest_dir(pattern: str) -> Optional[Path]:
    dirs = sorted([p for p in (PROJECT_ROOT / "results").glob(pattern) if p.is_dir()])
    return dirs[-1] if dirs else None


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_header_footer(doc: Document, title: str) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.text = title
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer)


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Prepared from the current DelayDispute Copilot repository, with the latest "
        "audit-ready results, forensic audit outputs, and manuscript support assets."
    )
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Times New Roman"


def add_paragraph(doc: Document, text: str, bold_prefix: Optional[str] = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        first.bold = True
        first.font.name = "Times New Roman"
        rest = p.add_run(text[len(bold_prefix):])
        rest.font.name = "Times New Roman"
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"


def clean_markdown_text(text: str) -> List[str]:
    lines: List[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        if line.startswith("#"):
            continue
        if line.startswith("- "):
            lines.append(line)
        else:
            lines.append(line)
    return lines


def insert_markdown_block(doc: Document, title: str, md_path: Path) -> None:
    if not md_path.exists():
        return
    add_heading(doc, title, 2)
    lines = clean_markdown_text(md_path.read_text(encoding="utf-8"))
    buffer: List[str] = []
    for line in lines:
        if not line:
            if buffer:
                add_paragraph(doc, " ".join(buffer))
                buffer = []
            continue
        if line.startswith("- "):
            if buffer:
                add_paragraph(doc, " ".join(buffer))
                buffer = []
            add_bullets(doc, [line[2:]])
        else:
            buffer.append(line)
    if buffer:
        add_paragraph(doc, " ".join(buffer))


def dataframe_to_word_table(doc: Document, df: pd.DataFrame, title: str, max_rows: Optional[int] = None) -> None:
    add_heading(doc, title, 2)
    if df.empty:
        add_paragraph(doc, "No rows were available for this table.")
        return
    if max_rows is not None:
        df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            if pd.isna(value):
                cells[i].text = ""
            elif isinstance(value, float):
                cells[i].text = f"{value:.4f}"
            else:
                cells[i].text = str(value)
    doc.add_paragraph("")


def add_figure(doc: Document, title: str, image_path: Path, width: float = 6.0) -> None:
    if not image_path.exists():
        return
    add_heading(doc, title, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"Source: {image_path.relative_to(PROJECT_ROOT).as_posix()}")
    run.italic = True
    run.font.name = "Times New Roman"


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--forensic_dir", type=str, default="")
    args = ap.parse_args()

    forensic_dir = (PROJECT_ROOT / args.forensic_dir).resolve() if args.forensic_dir else latest_dir("forensic_audit_*")
    if forensic_dir is None:
        raise FileNotFoundError("No forensic_audit_* directory found.")
    paper_assets = PROJECT_ROOT / "paper_assets"
    tables_dir = paper_assets / "tables"
    text_dir = paper_assets / "text"
    figures_dir = paper_assets / "figures"
    reports_dir = paper_assets / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    main_results = pd.read_csv(tables_dir / "table_main_result_claims.csv", encoding="utf-8-sig")
    baseline_perf = pd.read_csv(tables_dir / "table2_baseline_performance.csv", encoding="utf-8-sig")
    resp_metrics = pd.read_csv(tables_dir / "table4_responsibility_metrics.csv", encoding="utf-8-sig")
    audit_metrics = pd.read_csv(tables_dir / "table5_auditability_metrics.csv", encoding="utf-8-sig")
    dataset_profile = pd.read_csv(tables_dir / "table1_dataset_profile.csv", encoding="utf-8-sig")
    top5_actions = pd.read_csv(tables_dir / "table_top5_next_actions.csv", encoding="utf-8-sig")
    leakage_df = pd.read_csv(forensic_dir / "leakage_sentinel_results.csv", encoding="utf-8-sig")
    delta_df = pd.read_csv(forensic_dir / "delta_table.csv", encoding="utf-8-sig")
    claim_df = pd.read_csv(forensic_dir / "claim_tiering.csv", encoding="utf-8-sig")
    metric_recompute = pd.read_csv(forensic_dir / "metric_recompute_check.csv", encoding="utf-8-sig")
    resp_root = pd.read_csv(forensic_dir / "responsibility_root_cause.csv", encoding="utf-8-sig")

    title = "DelayDispute Copilot: Current Research Logic, Empirical Results, and Forensic Audit Synthesis"
    subtitle = f"Learning report generated on {datetime.now().strftime('%Y-%m-%d')}"

    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "DelayDispute Copilot Learning Report")
    add_title_page(doc, title, subtitle)

    add_heading(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "This report consolidates the current state of the DelayDispute Copilot project into a paper-style English document for study and manuscript preparation. "
        "The project is framed as a leakage-aware, audit-ready, management-oriented decision-support system for construction schedule delay disputes. "
        "Its present scope covers outcome prediction, responsibility diagnosis, auditable evidence-chain reconstruction, and forensic reproducibility auditing. "
        "The document integrates the current benchmark logic, empirical results, forensic audit conclusions, managerial relevance, limitations, and the most credible next actions for improving performance without weakening auditability."
    )

    add_heading(doc, "1. Study Positioning and Current Research Logic", 1)
    add_paragraph(
        doc,
        "The project is not positioned as a generic legal judgment classifier. It is framed as an engineering management decision-support framework that helps project teams anticipate delay-dispute outcomes, diagnose likely responsibility patterns, identify documentation and procedural weaknesses, and prepare for dispute triage before formal adjudication."
    )
    add_paragraph(
        doc,
        "The core scientific constraint is that model inputs must be restricted to pre-decision information. Post-decision text may be used only for candidate-label derivation, audit anchoring, and retrospective evaluation, but not for the main predictive pipeline. This boundary is necessary because the intended managerial use case is prospective dispute governance rather than retrospective case replay."
    )
    add_bullets(doc, [
        "Task 1: outcome prediction for delay-related claims.",
        "Task 2: structured responsibility diagnosis rather than vague one-line hints.",
        "Task 3: evidence-chain reconstruction with traceable spans and auditability checks.",
        "Task 4: forensic reproducibility auditing to define what can and cannot be claimed in the paper.",
    ])

    add_heading(doc, "2. Data Assets and Benchmark Governance", 1)
    add_paragraph(
        doc,
        "The study uses a weak-label pool of approximately 4,592 construction-delay dispute cases, together with a seed reference file and two machine-assisted candidate benchmarks. These benchmarks are not described as human gold. Instead, they are treated as candidate evaluation sets with explicit uncertainty and provenance tracking."
    )
    dataframe_to_word_table(doc, dataset_profile, "Table 1. Dataset and Label Distribution", max_rows=20)
    add_paragraph(
        doc,
        "The strict candidate benchmark contains 250 cases and the extended candidate benchmark contains 500 cases. Outcome labels are distributed across support, partial, and not_support. Responsibility labels are substantially less dominated by unknown than in the earlier gold500 file, making the current benchmark harder but more credible."
    )

    add_heading(doc, "3. Pipeline Logic and Auditability Design", 1)
    insert_markdown_block(doc, "3.1 Methodology Summary", text_dir / "methodology_summary.md")
    insert_markdown_block(doc, "3.2 Contribution Statement", text_dir / "contribution_statement.md")
    add_figure(doc, "Figure 1. Method Flow", figures_dir / "fig1_method_flow.png", width=6.2)
    add_paragraph(
        doc,
        "The current main method is the PAESC hybrid, which combines pre-decision text features, structured case signals, retrieval-style prior support, structured responsibility outputs, and evidence-chain auditing. The older current_hybrid_baseline remains important because it is the strongest reproducible predictive baseline on the present candidate benchmark."
    )

    add_heading(doc, "4. Main Empirical Results", 1)
    dataframe_to_word_table(doc, main_results, "Table 2. Main Result Claim Structure", max_rows=10)
    dataframe_to_word_table(doc, baseline_perf, "Table 3. Baseline and Main-Method Performance", max_rows=20)
    add_figure(doc, "Figure 2. Outcome Prediction Performance", figures_dir / "fig3_outcome_performance.png", width=6.2)
    add_paragraph(
        doc,
        "On the extended candidate benchmark, the audit-ready PAESC hybrid achieved an accuracy of 0.724 and a macro-F1 of 0.5658. "
        "The strongest reproducible baseline, current_hybrid_baseline, achieved an accuracy of 0.830 and a macro-F1 of 0.6731. "
        "This means the project currently has a clear trade-off: the more audit-ready and structurally interpretable method is not yet the highest-scoring method on outcome prediction."
    )
    add_paragraph(
        doc,
        "Traditional TF-IDF baselines remained substantially weaker, with macro-F1 values around 0.24 to 0.26 on the extended benchmark. "
        "Therefore, the project clearly exceeds simple text classification baselines, but the main technical challenge is now to recover predictive strength without sacrificing auditability."
    )

    add_heading(doc, "5. Responsibility Diagnosis and Evidence-Chain Results", 1)
    dataframe_to_word_table(doc, resp_metrics, "Table 4. Responsibility Diagnosis Performance", max_rows=20)
    dataframe_to_word_table(doc, audit_metrics, "Table 5. Evidence-Chain Auditability", max_rows=10)
    add_figure(doc, "Figure 3. Responsibility Performance", figures_dir / "fig5_responsibility_performance.png", width=6.0)
    add_figure(doc, "Figure 4. Evidence-Chain Auditability", figures_dir / "fig6_evidence_auditability.png", width=6.0)
    add_paragraph(
        doc,
        "Responsibility diagnosis is currently the weakest task. On the extended candidate benchmark, responsibility macro-F1 remains around 0.196. "
        "By contrast, evidence-chain auditability is strong: valid-span rate and pre-decision-span rate remain above 0.98. "
        "This indicates that the system is already strong at constructing traceable evidence packages, but still weaker at converting those packages into fine-grained responsibility labels."
    )

    add_heading(doc, "6. Forensic Audit: Why Earlier Scores Were Higher", 1)
    insert_markdown_block(doc, "6.1 Forensic Summary", text_dir / "forensic_summary.md")
    insert_markdown_block(doc, "6.2 Why the Old High Scores Cannot Be Used as the Main Claim", text_dir / "old_high_score_explanation.md")
    dataframe_to_word_table(doc, claim_df, "Table 6. Claim Tiering", max_rows=10)
    dataframe_to_word_table(doc, delta_df, "Table 7. Delta Table for Historical vs Current Runs", max_rows=13)
    add_paragraph(
        doc,
        "The forensic audit confirmed that the old high scores are real in the narrow sense that they can be recomputed from prediction-level artifacts. "
        "However, they are not valid headline results for the current paper because they were obtained under different evaluation regimes, weaker audit constraints, or much easier label distributions."
    )
    add_paragraph(
        doc,
        "The most important single finding is the old gold500 responsibility label distribution: approximately 82.6% of responsibility labels were unknown. "
        "This makes a very high responsibility accuracy or macro-F1 much easier to obtain than under the current candidate benchmarks, where unknown is only around 6.8% to 7.2%."
    )

    add_heading(doc, "7. Leakage Sentinel and Score Inflation Risk", 1)
    dataframe_to_word_table(doc, leakage_df, "Table 8. Leakage Sentinel Results", max_rows=10)
    add_paragraph(
        doc,
        "The leakage sentinel compares three settings: pre_decision only, post_decision only, and pre_decision plus post_decision. "
        "On the extended candidate benchmark, moving from pre_decision only to pre_decision plus post_decision raises outcome macro-F1 by about 0.0073. "
        "This does not fully explain the old 0.8007 legacy score, but it confirms that post-decision information can inflate performance and that leakage control must remain a hard methodological rule."
    )

    add_heading(doc, "8. Root-Cause Analysis of Low Responsibility Performance", 1)
    dataframe_to_word_table(doc, resp_root, "Table 9. Responsibility Root-Cause Evidence", max_rows=20)
    add_paragraph(
        doc,
        "The root-cause audit shows that low responsibility performance should not be interpreted as a purely model-side failure. "
        "Instead, three factors are interacting. First, the label problem remains serious because current candidate responsibility labels still disagree heavily with upstream LLM hints and remain imbalanced. "
        "Second, the task formulation itself is difficult: collapsing the seven-class schema into a coarser four-class formulation materially improves macro-F1. "
        "Third, there is still a model problem, but the current structured diagnosis remains better than a direct LLM-hint baseline and better than a majority baseline."
    )

    add_heading(doc, "9. Managerial Relevance", 1)
    insert_markdown_block(doc, "9.1 Managerial Relevance Statement", text_dir / "managerial_relevance_statement.md")
    insert_markdown_block(doc, "9.2 Validation Statement", text_dir / "validation_statement.md")
    insert_markdown_block(doc, "9.3 Limitations", text_dir / "limitations_statement.md")
    add_paragraph(
        doc,
        "From an IEEE-TEM perspective, the project is strongest when framed as managerial decision support with human oversight. "
        "Its practical value lies in triaging disputes, flagging evidence gaps, surfacing procedural vulnerabilities, and providing consistent review-ready summaries for project governance teams."
    )

    add_heading(doc, "10. Recommended Next Steps", 1)
    dataframe_to_word_table(doc, top5_actions, "Table 10. Ranked Next Actions", max_rows=10)
    add_paragraph(
        doc,
        "The highest-payoff next move is not a blind model rewrite. The evidence currently points toward a controlled recovery strategy: "
        "bring the strongest reproducible baseline signal back into the audit-ready pipeline, explicitly log the prior block, and ablate it carefully. "
        "For responsibility diagnosis, the fastest path to a stronger paper result is likely to use a coarser primary schema in the main text, while keeping the full fine-grained schema as an appendix or secondary analysis."
    )

    add_heading(doc, "11. File Map for Study and Reproduction", 1)
    add_bullets(doc, [
        f"Current forensic audit folder: {forensic_dir.relative_to(PROJECT_ROOT).as_posix()}",
        "Current audit-ready main run: results/final_eval_20260409_194025",
        "Main result claim table: paper_assets/tables/table_main_result_claims.csv",
        "Forensic delta table: results/forensic_audit_20260413_111946/delta_table.csv",
        "Leakage sentinel: results/forensic_audit_20260413_111946/leakage_sentinel_results.csv",
        "Responsibility root cause: results/forensic_audit_20260413_111946/responsibility_root_cause.csv",
        "Top next actions: paper_assets/tables/table_top5_next_actions.csv",
    ])

    report_md = reports_dir / f"DelayDispute_Copilot_Learning_Report_{datetime.now().strftime('%Y%m%d')}.md"
    report_docx = reports_dir / f"DelayDispute_Copilot_Learning_Report_{datetime.now().strftime('%Y%m%d')}.docx"
    md_text = "\n".join([
        f"# {title}",
        "",
        f"Generated on {datetime.now().strftime('%Y-%m-%d')}.",
        "",
        "This Markdown file is the sidecar source summary for the Word report.",
        "",
        "See the DOCX report for the full paper-style layout and inserted figures/tables.",
        "",
        f"Forensic audit source: {forensic_dir.relative_to(PROJECT_ROOT).as_posix()}",
        "Current audit-ready main run: results/final_eval_20260409_194025",
    ])
    report_md.write_text(md_text, encoding="utf-8")
    doc.save(report_docx)
    print(f"[DONE] {report_docx}")


if __name__ == "__main__":
    main()
