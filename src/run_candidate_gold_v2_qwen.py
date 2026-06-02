# -*- coding: utf-8 -*-
"""Build a Qwen-assisted temporary candidate_gold_v2 and rerun 500 cases.

This script is for workflow validation and plotting. It does not create human
gold labels. It uses post-decision text only for temporary label construction
and pre-decision text only for prediction.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

import pandas as pd
import requests
from tqdm import tqdm

try:
    import docx
except Exception as exc:  # pragma: no cover
    docx = None
    DOCX_ERROR = exc
else:
    DOCX_ERROR = None

try:
    import matplotlib.pyplot as plt
except Exception:  # pragma: no cover
    plt = None

try:
    import yaml
except Exception:  # pragma: no cover
    yaml = None

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.audit_utils import build_run_manifest, synthetic_file_diff, write_manifest, write_synthetic_git_summary  # noqa: E402
from src.run_true_llm_copilot import (  # noqa: E402
    FOLDED_RESP_LABELS,
    LABELS,
    RESP_LABELS,
    fold_responsibility,
    legacy_debug_key_from_script,
    recompute_outcome_metrics,
    rel,
    sha256_text,
)

SPLIT_ANCHORS = ["本院认为", "法院认为", "本院经审理认为", "本院审理认为", "裁判如下", "判决如下", "综上所述", "综上"]
ROLE_ALIASES = {
    "ENT": "entitlement",
    "NOT": "notice_substantiation",
    "CAU": "causality",
    "IMP": "impact_schedule_relevance",
    "DOC": "documentation_integrity",
    "entitlement": "entitlement",
    "notice_substantiation": "notice_substantiation",
    "causality": "causality",
    "impact_schedule_relevance": "impact_schedule_relevance",
    "documentation_integrity": "documentation_integrity",
}
RESP_MAP = {
    "业主": "owner",
    "发包人": "owner",
    "建设单位": "owner",
    "甲方": "owner",
    "owner": "owner",
    "承包人": "contractor",
    "承包商": "contractor",
    "施工单位": "contractor",
    "乙方": "contractor",
    "contractor": "contractor",
    "分包人": "subcontractor",
    "分包商": "subcontractor",
    "subcontractor": "subcontractor",
    "设计": "designer_supervisor",
    "监理": "designer_supervisor",
    "设计/监理": "designer_supervisor",
    "designer_supervisor": "designer_supervisor",
    "双方": "both",
    "共同责任": "both",
    "both": "both",
    "不可抗力": "force_majeure_policy",
    "政策": "force_majeure_policy",
    "force_majeure_policy": "force_majeure_policy",
    "不明确": "unknown",
    "无法判断": "unknown",
    "unknown": "unknown",
    "": "unknown",
}

DEFAULT_CFG = {
    "paths": {
        "raw_docx_dir": "data/0_raw_docx",
        "structured_case_dir": "data/3_structured_cases",
        "candidate_gold_strict_csv": "data/gold/candidate_gold_strict_v1.csv",
        "candidate_gold_extended_csv": "data/gold/candidate_gold_extended_v1.csv",
        "legacy_key_source_py": "src/llm_step2_fast_qc.py",
        "output_root": "results",
        "gold_output_dir": "data/gold",
        "figure_output_dir": "paper_assets/figures_candidate_gold_v2",
        "figure_data_dir": "paper_assets/figure_data_candidate_gold_v2",
        "manuscript_dir": "paper_assets/manuscript",
    },
    "llm": {
        "provider": "dashscope_qwen",
        "api_base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "api_key_env": "DASHSCOPE_API_KEY",
        "debug_api_key": "",
        "reuse_legacy_debug_key": False,
        "model_candidates": ["qwen-max", "qwen-plus"],
        "temperature": 0.0,
        "max_tokens_label": 1200,
        "max_tokens_prediction": 1600,
        "timeout": 120,
        "retries": 2,
        "workers": 3,
        "sleep_seconds": 0.05,
        "max_chars_pre": 9000,
        "max_chars_post": 9000,
        "prompt_template_version_label": "candidate_gold_v2_qwen_label_anchor_v1",
        "prompt_template_version_prediction": "candidate_gold_v2_qwen_predecision_predict_v1",
    },
    "eval": {"seed": 2026, "min_success_rate": 0.95},
    "run": {"run_name_prefix": "candidate_gold_v2_qwen"},
}


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def deep_update(base: Dict[str, Any], updates: Dict[str, Any]) -> Dict[str, Any]:
    for key, value in updates.items():
        if isinstance(value, dict) and isinstance(base.get(key), dict):
            deep_update(base[key], value)
        else:
            base[key] = value
    return base


def write_default_config(path: Path) -> None:
    if path.exists():
        return
    ensure_dir(path.parent)
    if yaml is not None:
        path.write_text(yaml.safe_dump(DEFAULT_CFG, allow_unicode=True, sort_keys=False), encoding="utf-8")
    else:
        path.write_text(json.dumps(DEFAULT_CFG, ensure_ascii=False, indent=2), encoding="utf-8")


def load_config(path: Path) -> Dict[str, Any]:
    cfg = json.loads(json.dumps(DEFAULT_CFG, ensure_ascii=False))
    if path.exists() and yaml is not None:
        cfg = deep_update(cfg, yaml.safe_load(path.read_text(encoding="utf-8")) or {})
    return cfg


def resolve_api_key(cfg: Dict[str, Any]) -> Tuple[str, str]:
    env_name = str(cfg["llm"].get("api_key_env", "DASHSCOPE_API_KEY"))
    key = os.getenv(env_name, "").strip()
    if key:
        return key, f"env:{env_name}"
    key = str(cfg["llm"].get("debug_api_key", "")).strip()
    if key:
        return key, "config:debug_api_key"
    if bool(cfg["llm"].get("reuse_legacy_debug_key", True)):
        p = PROJECT_ROOT / cfg["paths"].get("legacy_key_source_py", "src/llm_step2_fast_qc.py")
        key = legacy_debug_key_from_script(p)
        if key:
            return key, f"legacy_script:{rel(p)}"
    return "", "missing"


def normalize_label_v2(value: Any) -> str:
    raw = str(value or "").strip().lower()
    mapping = {
        "support": "support",
        "supported": "support",
        "支持": "support",
        "予以支持": "support",
        "全部支持": "support",
        "partial": "partial",
        "partial_support": "partial",
        "partially_support": "partial",
        "部分支持": "partial",
        "部分予以支持": "partial",
        "酌情支持": "partial",
        "not_support": "not_support",
        "not-support": "not_support",
        "reject": "not_support",
        "rejected": "not_support",
        "不支持": "not_support",
        "不予支持": "not_support",
        "驳回": "not_support",
        "全部驳回": "not_support",
        "unknown": "unknown",
        "nan": "unknown",
        "": "unknown",
    }
    if raw in mapping:
        return mapping[raw]
    if "部分" in raw or "酌情" in raw:
        return "partial"
    if "驳回" in raw or "不予" in raw or "不支持" in raw:
        return "not_support"
    if "支持" in raw:
        return "support"
    return "unknown"


def normalize_resp_v2(value: Any) -> str:
    raw = str(value or "").strip()
    if raw in RESP_LABELS:
        return raw
    if raw in RESP_MAP:
        return RESP_MAP[raw]
    for key, value2 in RESP_MAP.items():
        if key and key in raw:
            return value2
    return "unknown"


def extract_json_object(raw: str) -> Dict[str, Any]:
    text = str(raw or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    start, end = text.find("{"), text.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("No JSON object found in response")
    return json.loads(text[start : end + 1])


def split_pre_post_text(text: str) -> Tuple[str, str, str, float]:
    candidates = [(str(text).find(a), a) for a in SPLIT_ANCHORS if str(text).find(a) >= 0]
    if candidates:
        idx, anchor = sorted(candidates, key=lambda x: x[0])[0]
        return str(text)[:idx].strip(), str(text)[idx:].strip(), anchor, 0.9
    cut = int(len(str(text)) * 0.72)
    return str(text)[:cut].strip(), str(text)[cut:].strip(), "heuristic_72pct", 0.45


def compact_text(text: str, max_chars: int, keywords: Optional[Sequence[str]] = None) -> str:
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(text) <= max_chars:
        return text
    keywords = list(keywords or ["工期", "延期", "延误", "停工", "签证", "通知", "索赔", "关键线路", "证据", "违约", "判决", "驳回"])
    head = text[: int(max_chars * 0.35)]
    tail = text[-int(max_chars * 0.25) :]
    hits: List[str] = []
    for sent in re.split(r"(?<=[。！？；])", text):
        if any(k in sent for k in keywords):
            hits.append(sent.strip())
        if sum(len(x) for x in hits) >= int(max_chars * 0.35):
            break
    return (head + "\n【相关片段】\n" + "\n".join(hits) + "\n【尾部片段】\n" + tail)[:max_chars]


def read_docx_text(path: Path) -> str:
    if docx is None:
        raise RuntimeError(f"python-docx unavailable: {DOCX_ERROR}")
    document = docx.Document(str(path))
    chunks: List[str] = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def load_structured_fallback(case_id: str, structured_dir: Path) -> Dict[str, Any]:
    p = structured_dir / f"{case_id}.json"
    if not p.exists():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def build_case_text(row: pd.Series, cfg: Dict[str, Any]) -> Dict[str, Any]:
    cid = str(row["case_id"])
    docx_path = PROJECT_ROOT / cfg["paths"]["raw_docx_dir"] / str(row.get("source_file", ""))
    structured = load_structured_fallback(cid, PROJECT_ROOT / cfg["paths"]["structured_case_dir"])
    if docx_path.exists():
        text = read_docx_text(docx_path)
        source = "raw_docx"
    else:
        text = str(structured.get("raw_text", ""))
        source = "structured_fallback"
    pre, post, anchor, conf = split_pre_post_text(text)
    if not pre and structured.get("pre_decision_text"):
        pre = str(structured.get("pre_decision_text", ""))
    if not post and structured.get("post_decision_text"):
        post = str(structured.get("post_decision_text", ""))
    return {
        "case_id": cid,
        "source_file": str(row.get("source_file", "")),
        "docx_exists": int(docx_path.exists()),
        "text_source": source,
        "pre_decision_text": pre,
        "post_decision_text": post,
        "split_anchor": anchor,
        "pre_post_split_confidence_v2": conf,
    }


def role_name(value: Any) -> str:
    return ROLE_ALIASES.get(str(value or "").strip(), str(value or "").strip() or "unknown")


def validate_evidence_chain(chain: Any, pre_text: str) -> Tuple[List[Dict[str, Any]], Dict[str, float]]:
    rows: List[Dict[str, Any]] = []
    seen = set()
    if not isinstance(chain, list):
        chain = []
    for item in chain:
        if not isinstance(item, dict):
            continue
        span = str(item.get("span_text", item.get("text", "")) or "").strip()
        start = str(pre_text or "").find(span) if span else -1
        end = start + len(span) if start >= 0 else -1
        dup = int(bool(span) and span in seen)
        if span:
            seen.add(span)
        rows.append({
            "role_label": role_name(item.get("role_label", "")),
            "span_text": span,
            "span_start": start,
            "span_end": end,
            "pre_decision_flag": int(start >= 0),
            "valid_span_flag": int(start >= 0 and end <= len(str(pre_text or ""))),
            "duplicate_flag": dup,
            "reason": str(item.get("reason", ""))[:300],
        })
    expected = {"entitlement", "notice_substantiation", "causality", "impact_schedule_relevance", "documentation_integrity"}
    roles = {r["role_label"] for r in rows if r["valid_span_flag"] == 1}
    denom = max(1, len(rows))
    return rows, {
        "valid_span_rate": round(sum(r["valid_span_flag"] for r in rows) / denom, 6),
        "pre_decision_span_rate": round(sum(r["pre_decision_flag"] for r in rows) / denom, 6),
        "duplicate_chain_rate": round(sum(r["duplicate_flag"] for r in rows) / denom, 6),
        "role_coverage_rate": round(len(roles & expected) / len(expected), 6),
        "missing_role_rate": round(1 - len(roles & expected) / len(expected), 6),
    }


class QwenClient:
    def __init__(self, cfg: Dict[str, Any], api_key: str, model_name: str):
        self.base_url = str(cfg["llm"]["api_base_url"]).rstrip("/")
        self.model_name = model_name
        self.temperature = float(cfg["llm"].get("temperature", 0.0))
        self.timeout = int(cfg["llm"].get("timeout", 120))
        self.api_key = api_key

    def chat(self, messages: List[Dict[str, str]], max_tokens: int) -> Tuple[str, Dict[str, Any]]:
        payload = {"model": self.model_name, "messages": messages, "temperature": self.temperature, "max_tokens": int(max_tokens)}
        headers = {"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"}
        response = requests.post(f"{self.base_url}/chat/completions", headers=headers, json=payload, timeout=self.timeout)
        response.raise_for_status()
        obj = response.json()
        return obj["choices"][0]["message"]["content"], obj.get("usage", {})


def probe_model(cfg: Dict[str, Any], api_key: str, out_dir: Path) -> str:
    rows = []
    for model in cfg["llm"].get("model_candidates", ["qwen-max", "qwen-plus"]):
        started = time.time()
        try:
            content, usage = QwenClient(cfg, api_key, str(model)).chat(
                [{"role": "system", "content": "Return JSON only."}, {"role": "user", "content": "输出 {\"ok\":true}"}],
                max_tokens=80,
            )
            extract_json_object(content)
            rows.append({"model_name": model, "probe_status": "available", "latency_sec": round(time.time() - started, 4), "usage_json": json.dumps(usage, ensure_ascii=False)})
            pd.DataFrame(rows).to_csv(out_dir / "model_probe_manifest.csv", index=False, encoding="utf-8-sig")
            return str(model)
        except Exception as exc:
            rows.append({"model_name": model, "probe_status": "failed", "latency_sec": round(time.time() - started, 4), "error": str(exc)[:800]})
    pd.DataFrame(rows).to_csv(out_dir / "model_probe_manifest.csv", index=False, encoding="utf-8-sig")
    raise RuntimeError("No configured Qwen model is available")


def label_prompt(case: Dict[str, Any], row: pd.Series, cfg: Dict[str, Any]) -> Tuple[List[Dict[str, str]], str]:
    payload = {
        "case_id": case["case_id"],
        "task": "temporary_qwen_assisted_candidate_label_v2",
        "note": "This is a machine-assisted temporary candidate label, not human gold.",
        "schema": {"candidate_outcome_label_v2": LABELS, "candidate_responsibility_label_v2": RESP_LABELS},
        "rules": [
            "Use court reasoning and dispositive text as adjudication anchors.",
            "Prioritize delay damages, extension of time, liquidated damages for delay, stoppage, delay payment, causality, and responsibility.",
            "support=substantially supported; partial=partly supported or mixed; not_support=rejected or evidence-insufficient.",
            "Return exact anchor excerpts when possible.",
        ],
        "required_json": {
            "candidate_outcome_label_v2": "support|partial|not_support",
            "candidate_responsibility_label_v2": "|".join(RESP_LABELS),
            "outcome_anchor_text": "exact excerpt",
            "responsibility_anchor_text": "exact excerpt",
            "evidence_span_v2": "1-3 anchor excerpts separated by |",
            "confidence": "float 0-1",
            "conflict_flag": "boolean",
            "needs_review": "boolean",
            "note": "short Chinese audit note",
        },
        "old_machine_labels": {
            "candidate_outcome_label_v1": row.get("candidate_outcome_label", ""),
            "candidate_responsibility_label_v1": row.get("candidate_responsibility_label", ""),
            "weak_label": row.get("weak_label", ""),
            "llm_label": row.get("llm_label", ""),
            "post_label": row.get("post_label", ""),
        },
        "v1_evidence_span": str(row.get("evidence_span", ""))[:1600],
        "pre_decision_brief_for_context_only": compact_text(case["pre_decision_text"], 1800),
        "post_decision_text": compact_text(case["post_decision_text"], int(cfg["llm"].get("max_chars_post", 9000))),
    }
    prompt = json.dumps(payload, ensure_ascii=False)
    return [
        {"role": "system", "content": "你是建设工程工期延误纠纷标签审计助手。只输出一个合法 JSON 对象，不要 Markdown。"},
        {"role": "user", "content": prompt},
    ], prompt


def prediction_prompt(case: Dict[str, Any], cfg: Dict[str, Any]) -> Tuple[List[Dict[str, str]], str]:
    payload = {
        "case_id": case["case_id"],
        "task": "pre_decision_delay_dispute_prediction_v2",
        "hard_constraints": ["Use only pre_decision_text.", "Do not use judgment result or court reasoning.", "Return JSON only.", "span_text must be exact excerpts from pre_decision_text."],
        "schema": {"outcome_label": LABELS, "primary_responsible_party": RESP_LABELS, "evidence_roles": ["ENT", "NOT", "CAU", "IMP", "DOC"]},
        "required_json": {
            "outcome_label": "support|partial|not_support",
            "outcome_confidence": "float 0-1",
            "primary_responsible_party": "|".join(RESP_LABELS),
            "secondary_responsible_party": "|".join(RESP_LABELS),
            "responsibility_type": "single_party|shared|external|uncertain",
            "responsibility_confidence": "float 0-1",
            "uncertainty_flag": "boolean",
            "evidence_chain": [{"role_label": "ENT|NOT|CAU|IMP|DOC", "span_text": "exact excerpt", "reason": "short reason"}],
            "delay_irac": {"issue": "", "rule": "", "application": "", "conclusion": "", "management_action": ""},
            "documentation_gap_index": "float 0-1",
            "procedural_compliance_risk": "float 0-1",
            "causality_ambiguity": "float 0-1",
            "concurrency_risk": "float 0-1",
            "critical_path_support": "float 0-1",
            "negotiation_readiness_score": "float 0-1",
            "managerial_failure_type": "short label",
            "recommended_management_action": "short action",
        },
        "pre_decision_text": compact_text(case["pre_decision_text"], int(cfg["llm"].get("max_chars_pre", 9000))),
    }
    prompt = json.dumps(payload, ensure_ascii=False)
    return [
        {"role": "system", "content": "你是工程管理争议治理 Copilot。只能基于裁决前信息预测和生成证据链。只输出 JSON。"},
        {"role": "user", "content": prompt},
    ], prompt


def call_with_retries(client: QwenClient, messages: List[Dict[str, str]], max_tokens: int, retries: int) -> Tuple[str, Dict[str, Any], int]:
    last: Optional[Exception] = None
    for attempt in range(retries + 1):
        try:
            content, usage = client.chat(messages, max_tokens=max_tokens)
            return content, usage, attempt + 1
        except Exception as exc:
            last = exc
            if attempt < retries:
                time.sleep(1 + attempt)
    raise RuntimeError(str(last))


def clamp01(value: Any) -> float:
    try:
        return round(max(0.0, min(1.0, float(value))), 6)
    except Exception:
        return 0.0


def run_label_case(cid: str, case: Dict[str, Any], old_row: pd.Series, cfg: Dict[str, Any], client: QwenClient) -> Tuple[Dict[str, Any], Dict[str, Any], Dict[str, Any]]:
    messages, prompt = label_prompt(case, old_row, cfg)
    started = time.time()
    manifest = {"case_id": cid, "task_type": "label_v2", "model_name": client.model_name, "prompt_sha256": sha256_text(prompt), "prompt_chars": len(prompt)}
    raw = {"case_id": cid, "task_type": "label_v2", "prompt_sha256": manifest["prompt_sha256"], "raw_response": "", "usage_json": "{}"}
    try:
        content, usage, attempts = call_with_retries(client, messages, int(cfg["llm"].get("max_tokens_label", 1200)), int(cfg["llm"].get("retries", 2)))
        parsed = extract_json_object(content)
        label = normalize_label_v2(parsed.get("candidate_outcome_label_v2", "unknown"))
        resp = normalize_resp_v2(parsed.get("candidate_responsibility_label_v2", "unknown"))
        old_labels = [normalize_label_v2(old_row.get(c, "")) for c in ["candidate_outcome_label", "weak_label", "llm_label", "post_label"]]
        agreement = sum(1 for x in old_labels if x == label)
        rec = {
            "case_id": cid,
            "candidate_outcome_label_v2": label,
            "candidate_responsibility_label_v2": resp,
            "folded_responsibility_label_v2": fold_responsibility(resp),
            "outcome_anchor_text": str(parsed.get("outcome_anchor_text", ""))[:1000],
            "responsibility_anchor_text": str(parsed.get("responsibility_anchor_text", ""))[:1000],
            "evidence_span_v2": str(parsed.get("evidence_span_v2", ""))[:1800],
            "generation_source": f"qwen_assisted_label_v2::{client.model_name}",
            "confidence": clamp01(parsed.get("confidence", 0.0)),
            "conflict_flag": int(bool(parsed.get("conflict_flag", False)) or agreement < 2),
            "needs_review": int(bool(parsed.get("needs_review", False)) or agreement < 2),
            "source_agreement_count": agreement,
            "note": str(parsed.get("note", ""))[:500],
            "label_api_status": "api_available",
        }
        manifest.update({"api_status": "api_available", "parse_status": "parsed", "latency_sec": round(time.time() - started, 4), "attempts": attempts, "usage_json": json.dumps(usage, ensure_ascii=False)})
        raw.update({"raw_response": content, "usage_json": json.dumps(usage, ensure_ascii=False)})
        return rec, raw, manifest
    except Exception as exc:
        rec = {"case_id": cid, "candidate_outcome_label_v2": "unknown", "candidate_responsibility_label_v2": "unknown", "label_api_status": "api_error", "error": str(exc)[:1000]}
        manifest.update({"api_status": "api_error", "parse_status": "failed", "latency_sec": round(time.time() - started, 4), "error": str(exc)[:1000]})
        return rec, raw, manifest


def run_prediction_case(cid: str, case: Dict[str, Any], cfg: Dict[str, Any], client: QwenClient) -> Tuple[Dict[str, Any], Dict[str, Any], Dict[str, Any]]:
    messages, prompt = prediction_prompt(case, cfg)
    started = time.time()
    manifest = {"case_id": cid, "task_type": "prediction_v2", "model_name": client.model_name, "prompt_sha256": sha256_text(prompt), "prompt_chars": len(prompt)}
    raw = {"case_id": cid, "task_type": "prediction_v2", "prompt_sha256": manifest["prompt_sha256"], "raw_response": "", "usage_json": "{}"}
    try:
        content, usage, attempts = call_with_retries(client, messages, int(cfg["llm"].get("max_tokens_prediction", 1600)), int(cfg["llm"].get("retries", 2)))
        parsed = extract_json_object(content)
        chain, chain_metrics = validate_evidence_chain(parsed.get("evidence_chain", []), case["pre_decision_text"])
        irac = parsed.get("delay_irac", {}) if isinstance(parsed.get("delay_irac", {}), dict) else {}
        rec = {
            "case_id": cid,
            "prediction_api_status": "api_available",
            "outcome_label": normalize_label_v2(parsed.get("outcome_label", "unknown")),
            "outcome_confidence": clamp01(parsed.get("outcome_confidence", 0.0)),
            "primary_responsible_party": normalize_resp_v2(parsed.get("primary_responsible_party", "unknown")),
            "secondary_responsible_party": normalize_resp_v2(parsed.get("secondary_responsible_party", "unknown")),
            "responsibility_type": str(parsed.get("responsibility_type", "uncertain"))[:100],
            "responsibility_confidence": clamp01(parsed.get("responsibility_confidence", 0.0)),
            "uncertainty_flag": int(bool(parsed.get("uncertainty_flag", False))),
            "evidence_chain_json": json.dumps(chain, ensure_ascii=False),
            "delay_irac_json": json.dumps(irac, ensure_ascii=False),
            "issue": str(irac.get("issue", ""))[:300],
            "rule": str(irac.get("rule", ""))[:500],
            "application": str(irac.get("application", ""))[:800],
            "conclusion": str(irac.get("conclusion", ""))[:500],
            "management_action": str(irac.get("management_action", parsed.get("recommended_management_action", "")))[:500],
            "documentation_gap_index": clamp01(parsed.get("documentation_gap_index", 0.0)),
            "procedural_compliance_risk": clamp01(parsed.get("procedural_compliance_risk", 0.0)),
            "causality_ambiguity": clamp01(parsed.get("causality_ambiguity", 0.0)),
            "concurrency_risk": clamp01(parsed.get("concurrency_risk", 0.0)),
            "critical_path_support": clamp01(parsed.get("critical_path_support", 0.0)),
            "negotiation_readiness_score": clamp01(parsed.get("negotiation_readiness_score", 0.0)),
            "managerial_failure_type": str(parsed.get("managerial_failure_type", ""))[:120],
            "recommended_management_action": str(parsed.get("recommended_management_action", ""))[:500],
            **chain_metrics,
        }
        manifest.update({"api_status": "api_available", "parse_status": "parsed", "latency_sec": round(time.time() - started, 4), "attempts": attempts, "usage_json": json.dumps(usage, ensure_ascii=False)})
        raw.update({"raw_response": content, "usage_json": json.dumps(usage, ensure_ascii=False)})
        return rec, raw, manifest
    except Exception as exc:
        rec = {"case_id": cid, "prediction_api_status": "api_error", "outcome_label": "unknown", "error": str(exc)[:1000]}
        manifest.update({"api_status": "api_error", "parse_status": "failed", "latency_sec": round(time.time() - started, 4), "error": str(exc)[:1000]})
        return rec, raw, manifest


def run_batch(kind: str, ids: List[str], cases: Dict[str, Dict[str, Any]], old_rows: Dict[str, pd.Series], cfg: Dict[str, Any], client: QwenClient, out_dir: Path) -> pd.DataFrame:
    records_path = out_dir / f"{kind}_records.csv"
    manifest_path = out_dir / f"{kind}_api_manifest.csv"
    raw_path = out_dir / f"{kind}_raw_responses.jsonl"
    records: List[Dict[str, Any]] = []
    manifests: List[Dict[str, Any]] = []
    existing_ids = set()
    if records_path.exists():
        old = pd.read_csv(records_path, encoding="utf-8-sig")
        if "case_id" in old.columns:
            records = old.to_dict("records")
            existing_ids = set(old["case_id"].astype(str).tolist())
    if manifest_path.exists():
        old_manifest = pd.read_csv(manifest_path, encoding="utf-8-sig")
        manifests = old_manifest.to_dict("records")

    raws: List[Dict[str, Any]] = []
    pending_ids = [cid for cid in ids if cid not in existing_ids]
    if not pending_ids:
        return pd.DataFrame(records)

    workers = int(cfg["llm"].get("workers", 3))
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = []
        for cid in pending_ids:
            if kind == "label_v2":
                futures.append(executor.submit(run_label_case, cid, cases[cid], old_rows[cid], cfg, client))
            else:
                futures.append(executor.submit(run_prediction_case, cid, cases[cid], cfg, client))
            time.sleep(float(cfg["llm"].get("sleep_seconds", 0.05)))
        for idx, fut in enumerate(tqdm(as_completed(futures), total=len(futures), desc=f"Qwen {kind}"), start=1):
            rec, raw, manifest = fut.result()
            records.append(rec)
            raws.append(raw)
            manifests.append(manifest)
            if idx % 10 == 0:
                pd.DataFrame(records).to_csv(records_path, index=False, encoding="utf-8-sig")
                pd.DataFrame(manifests).to_csv(manifest_path, index=False, encoding="utf-8-sig")
                with raw_path.open("a", encoding="utf-8") as f:
                    for row in raws:
                        f.write(json.dumps(row, ensure_ascii=False) + "\n")
                raws = []
    pd.DataFrame(records).to_csv(records_path, index=False, encoding="utf-8-sig")
    pd.DataFrame(manifests).to_csv(manifest_path, index=False, encoding="utf-8-sig")
    with raw_path.open("a", encoding="utf-8") as f:
        for row in raws:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")
    return pd.DataFrame(records)


def load_candidate_rows(cfg: Dict[str, Any], max_cases: int) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    strict = pd.read_csv(PROJECT_ROOT / cfg["paths"]["candidate_gold_strict_csv"], encoding="utf-8-sig")
    extended = pd.read_csv(PROJECT_ROOT / cfg["paths"]["candidate_gold_extended_csv"], encoding="utf-8-sig")
    strict["dataset_name"] = "candidate_gold_strict_v2"
    extended["dataset_name"] = "candidate_gold_extended_v2"
    all_rows = pd.concat([strict, extended], ignore_index=True)
    all_rows["case_id"] = all_rows["case_id"].astype(str)
    dedup = all_rows.drop_duplicates("case_id", keep="last").copy()
    if max_cases > 0:
        keep = dedup["case_id"].tolist()[:max_cases]
        dedup = dedup[dedup["case_id"].isin(keep)].copy()
        strict = strict[strict["case_id"].astype(str).isin(keep)].copy()
        extended = extended[extended["case_id"].astype(str).isin(keep)].copy()
    return strict, extended, dedup


def write_candidate_files(label_df: pd.DataFrame, strict_v1: pd.DataFrame, extended_v1: pd.DataFrame, cfg: Dict[str, Any], out_dir: Path) -> Tuple[pd.DataFrame, pd.DataFrame]:
    gold_dir = ensure_dir(PROJECT_ROOT / cfg["paths"]["gold_output_dir"])
    for col in ["case_id", "candidate_outcome_label_v2", "candidate_responsibility_label_v2", "folded_responsibility_label_v2", "outcome_anchor_text", "responsibility_anchor_text", "evidence_span_v2", "generation_source", "confidence", "conflict_flag", "needs_review", "source_agreement_count", "note", "label_api_status"]:
        if col not in label_df.columns:
            label_df[col] = ""
    label_df.to_csv(gold_dir / "candidate_gold_v2_qwen.csv", index=False, encoding="utf-8-sig")

    def merge(base: pd.DataFrame, name: str) -> pd.DataFrame:
        m = base.merge(label_df, on="case_id", how="left", suffixes=("", "_label"))
        m["candidate_outcome_label_v1"] = m["candidate_outcome_label"].map(normalize_label_v2)
        m["candidate_responsibility_label_v1"] = m["candidate_responsibility_label"].map(normalize_resp_v2)
        m["candidate_outcome_label"] = m["candidate_outcome_label_v2"].map(normalize_label_v2)
        m["candidate_responsibility_label"] = m["candidate_responsibility_label_v2"].map(normalize_resp_v2)
        m["evidence_span"] = m["evidence_span_v2"].fillna(m.get("evidence_span", ""))
        m["generation_source"] = m["generation_source"].fillna("qwen_assisted_label_v2")
        m["dataset_name"] = name
        m.to_csv(gold_dir / f"{name}.csv", index=False, encoding="utf-8-sig")
        return m

    strict_v2 = merge(strict_v1, "candidate_gold_strict_v2")
    extended_v2 = merge(extended_v1, "candidate_gold_extended_v2")
    label_df.to_csv(out_dir / "candidate_gold_v2_records.csv", index=False, encoding="utf-8-sig")
    rows = []
    for name, df in [("candidate_gold_strict_v2", strict_v2), ("candidate_gold_extended_v2", extended_v2)]:
        valid = df[df["candidate_outcome_label"].isin(LABELS)]
        rows.append({
            "dataset_name": name,
            "n": int(len(valid)),
            "outcome_changed_rate": float((valid["candidate_outcome_label_v1"] != valid["candidate_outcome_label"]).mean()) if len(valid) else 0.0,
            "mean_label_confidence": float(pd.to_numeric(valid["confidence"], errors="coerce").fillna(0).mean()) if len(valid) else 0.0,
            "needs_review_rate": float(pd.to_numeric(valid["needs_review"], errors="coerce").fillna(1).mean()) if len(valid) else 0.0,
            "conflict_flag_rate": float(pd.to_numeric(valid["conflict_flag"], errors="coerce").fillna(1).mean()) if len(valid) else 0.0,
        })
    pd.DataFrame(rows).to_csv(out_dir / "label_change_summary.csv", index=False, encoding="utf-8-sig")
    try:
        with pd.ExcelWriter(gold_dir / "candidate_gold_v2_qwen_audit.xlsx") as writer:
            label_df.to_excel(writer, sheet_name="unique_500", index=False)
            strict_v2.to_excel(writer, sheet_name="strict_v2", index=False)
            extended_v2.to_excel(writer, sheet_name="extended_v2", index=False)
            pd.DataFrame(rows).to_excel(writer, sheet_name="change_summary", index=False)
    except Exception:
        pass
    return strict_v2, extended_v2


def write_eval_artifacts(pred_df: pd.DataFrame, strict_v2: pd.DataFrame, extended_v2: pd.DataFrame, out_dir: Path) -> Dict[str, Any]:
    eval_rows = pd.concat([strict_v2, extended_v2], ignore_index=True)
    merged = eval_rows.merge(pred_df, on="case_id", how="left")
    # Candidate files inherited several audit columns from v1. After merging,
    # prefer the fresh prediction-side columns so v2 figures are not populated
    # with stale or suffixed values.
    for col in [
        "valid_span_rate",
        "pre_decision_span_rate",
        "duplicate_chain_rate",
        "role_coverage_rate",
        "missing_role_rate",
        "documentation_gap_index",
        "procedural_compliance_risk",
        "causality_ambiguity",
        "concurrency_risk",
        "critical_path_support",
        "negotiation_readiness_score",
    ]:
        if f"{col}_y" in merged.columns:
            merged[col] = merged[f"{col}_y"]
        elif col not in merged.columns and f"{col}_x" in merged.columns:
            merged[col] = merged[f"{col}_x"]
    merged["model_name"] = "qwen_strong_direct_v2"
    merged["eval_split"] = "external_candidate_v2_eval"
    merged["y_true"] = merged["candidate_outcome_label"].map(normalize_label_v2)
    merged["y_pred"] = merged["outcome_label"].map(normalize_label_v2)
    merged["api_status"] = merged.get("prediction_api_status", "missing_or_failed").fillna("missing_or_failed")
    pred_cols = ["case_id", "dataset_name", "eval_split", "model_name", "y_true", "y_pred", "outcome_confidence", "candidate_outcome_label", "candidate_responsibility_label", "primary_responsible_party", "secondary_responsible_party", "responsibility_type", "responsibility_confidence", "uncertainty_flag", "api_status", "valid_span_rate", "pre_decision_span_rate", "role_coverage_rate", "documentation_gap_index", "procedural_compliance_risk", "causality_ambiguity", "concurrency_risk", "critical_path_support", "negotiation_readiness_score", "managerial_failure_type", "recommended_management_action"]
    for col in pred_cols:
        if col not in merged.columns:
            merged[col] = ""
    merged[pred_cols].to_csv(out_dir / "predictions_main.csv", index=False, encoding="utf-8-sig")
    valid = merged[merged["api_status"].eq("api_available") & merged["y_true"].isin(LABELS) & merged["y_pred"].isin(LABELS)]
    metrics: Dict[str, Any] = {"note": "Qwen-assisted temporary candidate labels, not human gold.", "candidate_gold_v2_evaluation": {}}
    per_rows, cm_rows = [], []
    for dataset, sub in valid.groupby("dataset_name"):
        m = recompute_outcome_metrics(sub[["y_true", "y_pred"]].to_dict("records"), LABELS)
        m["n_eval_rows"] = int(len(sub))
        metrics["candidate_gold_v2_evaluation"][dataset] = {"qwen_strong_direct_v2": m}
        for label in LABELS:
            item = m["per_class"].get(label, {})
            per_rows.append({"dataset_name": dataset, "model_name": "qwen_strong_direct_v2", "class_label": label, "precision": item.get("precision", 0), "recall": item.get("recall", 0), "f1": item.get("f1-score", 0), "support": item.get("support", 0)})
        for i, gold in enumerate(LABELS):
            for j, pred in enumerate(LABELS):
                cm_rows.append({"dataset_name": dataset, "model_name": "qwen_strong_direct_v2", "gold_label": gold, "pred_label": pred, "count": m["confusion_matrix"][i][j]})
    (out_dir / "metrics_main.json").write_text(json.dumps(metrics, ensure_ascii=False, indent=2), encoding="utf-8")
    pd.DataFrame(per_rows).to_csv(out_dir / "per_class_results.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame(cm_rows).to_csv(out_dir / "confusion_matrix_data.csv", index=False, encoding="utf-8-sig")

    resp_rows, chain_rows, mech_rows, err_rows = [], [], [], []
    for _, row in merged.iterrows():
        primary = normalize_resp_v2(row.get("primary_responsible_party", "unknown"))
        gold_resp = normalize_resp_v2(row.get("candidate_responsibility_label", "unknown"))
        resp_rows.append({"case_id": row["case_id"], "dataset_name": row["dataset_name"], "candidate_responsibility_label": gold_resp, "primary_responsible_party": primary, "folded_gold": fold_responsibility(gold_resp), "folded_pred": fold_responsibility(primary), "responsibility_confidence": row.get("responsibility_confidence", 0), "uncertainty_flag": row.get("uncertainty_flag", 1), "api_status": row.get("api_status")})
        chain_rows.append({"case_id": row["case_id"], "dataset_name": row["dataset_name"], "valid_span_rate": row.get("valid_span_rate", 0), "pre_decision_span_rate": row.get("pre_decision_span_rate", 0), "duplicate_chain_rate": row.get("duplicate_chain_rate", 0), "role_coverage_rate": row.get("role_coverage_rate", 0), "api_status": row.get("api_status")})
        mech_rows.append({"case_id": row["case_id"], "dataset_name": row["dataset_name"], "documentation_gap_index": row.get("documentation_gap_index", 0), "procedural_compliance_risk": row.get("procedural_compliance_risk", 0), "causality_ambiguity": row.get("causality_ambiguity", 0), "concurrency_risk": row.get("concurrency_risk", 0), "critical_path_support": row.get("critical_path_support", 0), "negotiation_readiness_score": row.get("negotiation_readiness_score", 0), "managerial_failure_type": row.get("managerial_failure_type", ""), "recommended_management_action": row.get("recommended_management_action", "")})
        if row.get("api_status") == "api_available" and row.get("y_true") in LABELS and row.get("y_pred") in LABELS and row.get("y_true") != row.get("y_pred"):
            err_rows.append({"case_id": row["case_id"], "dataset_name": row["dataset_name"], "y_true": row["y_true"], "y_pred": row["y_pred"], "error_category": "candidate_v2_label_prediction_mismatch"})
    pd.DataFrame(resp_rows).to_csv(out_dir / "responsibility_eval.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame(chain_rows).to_csv(out_dir / "evidence_chain_eval.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame(mech_rows).to_csv(out_dir / "managerial_mechanisms.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame(err_rows).to_csv(out_dir / "error_analysis.csv", index=False, encoding="utf-8-sig")
    summaries = []
    rdf = pd.DataFrame(resp_rows)
    for dataset, sub in rdf[rdf["api_status"].eq("api_available")].groupby("dataset_name"):
        fine = recompute_outcome_metrics([{"y_true": r["candidate_responsibility_label"], "y_pred": r["primary_responsible_party"]} for _, r in sub.iterrows()], RESP_LABELS)
        folded = recompute_outcome_metrics([{"y_true": r["folded_gold"], "y_pred": r["folded_pred"]} for _, r in sub.iterrows()], FOLDED_RESP_LABELS)
        summaries.append({"dataset_name": dataset, "fine_accuracy": fine["accuracy"], "fine_macro_f1": fine["macro_f1"], "folded_accuracy": folded["accuracy"], "folded_macro_f1": folded["macro_f1"]})
    pd.DataFrame(summaries).to_csv(out_dir / "responsibility_summary.csv", index=False, encoding="utf-8-sig")
    return metrics


def write_model_comparison(strict_v2: pd.DataFrame, extended_v2: pd.DataFrame, out_dir: Path) -> pd.DataFrame:
    rows = []
    ref_path = PROJECT_ROOT / "results" / "final_eval_20260409_194025" / "predictions_main.csv"
    eval_rows = pd.concat([strict_v2, extended_v2], ignore_index=True)[["case_id", "dataset_name", "candidate_outcome_label"]].copy()
    eval_rows["y_true_v2"] = eval_rows["candidate_outcome_label"].map(normalize_label_v2)
    if ref_path.exists():
        ref = pd.read_csv(ref_path, encoding="utf-8-sig")
        ref = ref[ref["model_name"].isin(["current_hybrid_baseline", "paesc_hybrid"])].copy()
        ref["dataset_name"] = ref["dataset_name"].replace({"candidate_gold_strict_v1": "candidate_gold_strict_v2", "candidate_gold_extended_v1": "candidate_gold_extended_v2"})
        joined = eval_rows.merge(ref[["case_id", "dataset_name", "model_name", "y_pred"]], on=["case_id", "dataset_name"], how="inner")
        for (dataset, model), sub in joined.groupby(["dataset_name", "model_name"]):
            m = recompute_outcome_metrics([{"y_true": r["y_true_v2"], "y_pred": normalize_label_v2(r["y_pred"])} for _, r in sub.iterrows()], LABELS)
            rows.append({"dataset_name": dataset, "model_name": model, "accuracy": m["accuracy"], "macro_f1": m["macro_f1"], "weighted_f1": m["weighted_f1"], "n": int(len(sub)), "comparison_note": "old_predictions_recomputed_against_candidate_gold_v2"})
    pred = pd.read_csv(out_dir / "predictions_main.csv", encoding="utf-8-sig")
    for (dataset, model), sub in pred.groupby(["dataset_name", "model_name"]):
        sub = sub[sub["y_true"].isin(LABELS) & sub["y_pred"].isin(LABELS)]
        m = recompute_outcome_metrics(sub[["y_true", "y_pred"]].to_dict("records"), LABELS)
        rows.append({"dataset_name": dataset, "model_name": model, "accuracy": m["accuracy"], "macro_f1": m["macro_f1"], "weighted_f1": m["weighted_f1"], "n": int(len(sub)), "comparison_note": "qwen_prediction_against_candidate_gold_v2"})
    comp = pd.DataFrame(rows)
    comp.to_csv(out_dir / "model_comparison_v2.csv", index=False, encoding="utf-8-sig")
    return comp


def infer_existing_model(out_dir: Path) -> str:
    for name in ["model_probe_manifest.csv", "label_v2_api_manifest.csv", "prediction_v2_api_manifest.csv"]:
        path = out_dir / name
        if not path.exists():
            continue
        try:
            df = pd.read_csv(path, encoding="utf-8-sig")
        except Exception:
            continue
        if "model_name" in df.columns and not df.empty:
            vals = [str(v) for v in df["model_name"].dropna().tolist() if str(v).strip()]
            if vals:
                return vals[0]
    return "records_only"


def write_figures(out_dir: Path, cfg: Dict[str, Any]) -> None:
    if plt is None:
        return
    fig_dir = ensure_dir(PROJECT_ROOT / cfg["paths"]["figure_output_dir"])
    data_dir = ensure_dir(PROJECT_ROOT / cfg["paths"]["figure_data_dir"])
    plt.rcParams.update({"font.family": ["Microsoft YaHei", "SimHei", "Arial"], "axes.unicode_minus": False, "figure.dpi": 160})
    comp = pd.read_csv(out_dir / "model_comparison_v2.csv", encoding="utf-8-sig")
    comp.to_csv(data_dir / "fig_v2_model_macro_f1.csv", index=False, encoding="utf-8-sig")
    try:
        comp.to_excel(data_dir / "fig_v2_model_macro_f1.xlsx", index=False)
    except Exception:
        pass
    pivot = comp.pivot_table(index="model_name", columns="dataset_name", values="macro_f1", aggfunc="first").fillna(0)
    ax = pivot.plot(kind="barh", figsize=(8, 4.4), color=["#4C78A8", "#72B7B2"])
    ax.set_xlabel("Macro-F1 against Qwen-assisted candidate_gold_v2")
    ax.set_ylabel("")
    ax.set_xlim(0, max(0.1, float(pivot.max().max()) * 1.15))
    ax.grid(axis="x", alpha=0.25)
    ax.legend(frameon=False, loc="lower right")
    plt.tight_layout()
    for ext in ["png", "pdf", "svg"]:
        plt.savefig(fig_dir / f"fig_v2_model_macro_f1.{ext}", bbox_inches="tight")
    plt.close()

    lab = pd.read_csv(out_dir / "label_change_summary.csv", encoding="utf-8-sig")
    lab.to_csv(data_dir / "fig_v2_label_change_summary.csv", index=False, encoding="utf-8-sig")
    try:
        lab.to_excel(data_dir / "fig_v2_label_change_summary.xlsx", index=False)
    except Exception:
        pass
    ax = lab.set_index("dataset_name")[["outcome_changed_rate", "needs_review_rate", "conflict_flag_rate"]].plot(kind="bar", figsize=(8, 4.2), color=["#F58518", "#B279A2", "#E45756"])
    ax.set_ylim(0, 1)
    ax.set_ylabel("Rate")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, ncol=3, fontsize=8)
    plt.xticks(rotation=0)
    plt.tight_layout()
    for ext in ["png", "pdf", "svg"]:
        plt.savefig(fig_dir / f"fig_v2_label_audit.{ext}", bbox_inches="tight")
    plt.close()

    chain = pd.read_csv(out_dir / "evidence_chain_eval.csv", encoding="utf-8-sig")
    summary = chain.groupby("dataset_name")[["valid_span_rate", "pre_decision_span_rate", "role_coverage_rate"]].mean().reset_index()
    summary.to_csv(data_dir / "fig_v2_evidence_auditability.csv", index=False, encoding="utf-8-sig")
    try:
        summary.to_excel(data_dir / "fig_v2_evidence_auditability.xlsx", index=False)
    except Exception:
        pass
    ax = summary.set_index("dataset_name").plot(kind="bar", figsize=(8, 4.2), color=["#54A24B", "#4C78A8", "#72B7B2"])
    ax.set_ylim(0, 1)
    ax.set_ylabel("Mean rate")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, ncol=3, fontsize=8)
    plt.xticks(rotation=0)
    plt.tight_layout()
    for ext in ["png", "pdf", "svg"]:
        plt.savefig(fig_dir / f"fig_v2_evidence_auditability.{ext}", bbox_inches="tight")
    plt.close()


def write_word_report(out_dir: Path, cfg: Dict[str, Any], model: str) -> Optional[Path]:
    if docx is None:
        return None
    report_dir = ensure_dir(PROJECT_ROOT / cfg["paths"]["manuscript_dir"])
    report = report_dir / f"candidate_gold_v2_qwen_rerun_report_{out_dir.name.replace('candidate_gold_v2_qwen_', '')}.docx"
    d = docx.Document()
    d.add_heading("Qwen-assisted candidate_gold_v2 rerun report", 0)
    d.add_paragraph("This report is for running through the 500-case workflow and inspecting plots. candidate_gold_v2 is Qwen-assisted and adjudication-anchored, not human gold. It should be replaced by the user's real manually reviewed 500 labels later.")
    d.add_paragraph(f"Run directory: {out_dir}")
    d.add_paragraph(f"Selected model: {model}")
    d.add_heading("Main metrics", level=1)
    metrics = json.loads((out_dir / "metrics_main.json").read_text(encoding="utf-8"))
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Dataset", "Model", "Accuracy", "Macro-F1", "Weighted-F1"]):
        table.rows[0].cells[i].text = h
    for dataset, obj in metrics["candidate_gold_v2_evaluation"].items():
        for mname, m in obj.items():
            cells = table.add_row().cells
            cells[0].text = dataset
            cells[1].text = mname
            cells[2].text = f"{m.get('accuracy', 0):.4f}"
            cells[3].text = f"{m.get('macro_f1', 0):.4f}"
            cells[4].text = f"{m.get('weighted_f1', 0):.4f}"
    d.add_heading("Model comparison", level=1)
    comp = pd.read_csv(out_dir / "model_comparison_v2.csv", encoding="utf-8-sig")
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Dataset", "Model", "Accuracy", "Macro-F1", "Note"]):
        table.rows[0].cells[i].text = h
    for _, r in comp.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["dataset_name"])
        cells[1].text = str(r["model_name"])
        cells[2].text = f"{float(r['accuracy']):.4f}"
        cells[3].text = f"{float(r['macro_f1']):.4f}"
        cells[4].text = str(r.get("comparison_note", ""))[:70]
    d.add_heading("Label audit", level=1)
    lab = pd.read_csv(out_dir / "label_change_summary.csv", encoding="utf-8-sig")
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Dataset", "N", "Outcome changed", "Needs review", "Mean confidence"]):
        table.rows[0].cells[i].text = h
    for _, r in lab.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["dataset_name"])
        cells[1].text = str(int(r["n"]))
        cells[2].text = f"{float(r['outcome_changed_rate']):.3f}"
        cells[3].text = f"{float(r['needs_review_rate']):.3f}"
        cells[4].text = f"{float(r['mean_label_confidence']):.3f}"
    d.add_heading("Figures", level=1)
    fig_dir = PROJECT_ROOT / cfg["paths"]["figure_output_dir"]
    for filename, caption in [
        ("fig_v2_model_macro_f1.png", "Model comparison under Qwen-assisted candidate_gold_v2 labels."),
        ("fig_v2_label_audit.png", "Label change, conflict, and review-need audit."),
        ("fig_v2_evidence_auditability.png", "Evidence-chain auditability in the rerun."),
    ]:
        p = fig_dir / filename
        if p.exists():
            d.add_picture(str(p), width=docx.shared.Inches(6.2))
            d.add_paragraph(caption)
    d.add_heading("Interpretation", level=1)
    d.add_paragraph("If F1 improves under candidate_gold_v2, the safe interpretation is that Qwen-assisted label alignment reduces temporary benchmark noise. It does not prove performance against human labels. Once the real manually reviewed 500 labels are available, replace the v2 files and rerun the same evaluation from prediction-level artifacts.")
    d.save(report)
    return report


def write_manifest_files(out_dir: Path, cfg_path: Path, cfg: Dict[str, Any], model: str, api_key_source: str, audit_status: str) -> None:
    artifacts = [
        cfg_path,
        PROJECT_ROOT / "src" / "run_candidate_gold_v2_qwen.py",
        PROJECT_ROOT / cfg["paths"]["candidate_gold_strict_csv"],
        PROJECT_ROOT / cfg["paths"]["candidate_gold_extended_csv"],
        PROJECT_ROOT / "data/gold/candidate_gold_strict_v2.csv",
        PROJECT_ROOT / "data/gold/candidate_gold_extended_v2.csv",
        out_dir / "predictions_main.csv",
        out_dir / "metrics_main.json",
        out_dir / "model_comparison_v2.csv",
        out_dir / "label_v2_api_manifest.csv",
        out_dir / "prediction_v2_api_manifest.csv",
    ]
    manifest = build_run_manifest(
        PROJECT_ROOT,
        PROJECT_ROOT / "requirements.txt",
        artifacts,
        model_name=model,
        prompt_template_version=f"{cfg['llm'].get('prompt_template_version_label')} + {cfg['llm'].get('prompt_template_version_prediction')}",
        embedding_model="none",
        label_schema_version="outcome_v1__responsibility_v1__candidate_gold_v2_qwen_temporary",
        command=" ".join(sys.argv),
        seed=int(cfg["eval"].get("seed", 2026)),
        split_mode="external_candidate_v2_eval",
        text_mode="label_from_post_decision__prediction_from_pre_decision_only",
        train_label_file=None,
        eval_label_file=PROJECT_ROOT / "data/gold/candidate_gold_extended_v2.csv",
        metric_source_files=[out_dir / "predictions_main.csv", out_dir / "responsibility_eval.csv", out_dir / "evidence_chain_eval.csv"],
        audit_status=audit_status,
        extra={"provider": cfg["llm"].get("provider"), "api_base_url": cfg["llm"].get("api_base_url"), "api_key_source": api_key_source, "candidate_gold_v2_note": "qwen-assisted temporary benchmark, not human gold"},
    )
    write_manifest(out_dir / "run_manifest.json", manifest)
    write_synthetic_git_summary(out_dir / "git_diff_summary.txt", "git unavailable in workspace; synthetic diff used")
    synthetic_file_diff([], [p for p in artifacts if p.exists()], PROJECT_ROOT, out_dir.name).to_csv(out_dir / "file_diff_summary.csv", index=False, encoding="utf-8-sig")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--config", default="config/research_candidate_gold_v2_qwen.yaml")
    parser.add_argument("--out-dir", default="")
    parser.add_argument("--max-cases", type=int, default=0)
    parser.add_argument("--workers", type=int, default=0)
    parser.add_argument("--skip-api", action="store_true")
    args = parser.parse_args()

    cfg_path = PROJECT_ROOT / args.config
    write_default_config(cfg_path)
    cfg = load_config(cfg_path)
    if args.workers > 0:
        cfg["llm"]["workers"] = args.workers
    out_dir = PROJECT_ROOT / args.out_dir if args.out_dir else PROJECT_ROOT / cfg["paths"]["output_root"] / f"{cfg['run']['run_name_prefix']}_{now_stamp()}"
    ensure_dir(out_dir)

    strict_v1, extended_v1, dedup = load_candidate_rows(cfg, args.max_cases)
    cases, old_rows = {}, {}
    for _, row in dedup.iterrows():
        cid = str(row["case_id"])
        cases[cid] = build_case_text(row, cfg)
        old_rows[cid] = row
    ids = list(cases.keys())
    print(f"Output dir: {out_dir}")
    print(f"Unique cases: {len(ids)}; strict rows: {len(strict_v1)}; extended rows: {len(extended_v1)}")

    model = "records_only"
    api_key, api_key_source = resolve_api_key(cfg)
    if args.skip_api:
        label_df = pd.read_csv(out_dir / "label_v2_records.csv", encoding="utf-8-sig")
        pred_df = pd.read_csv(out_dir / "prediction_v2_records.csv", encoding="utf-8-sig")
        model = infer_existing_model(out_dir)
    else:
        if not api_key:
            (out_dir / "api_unavailable.txt").write_text("No DashScope API key found.\n", encoding="utf-8")
            print("No DashScope API key found.")
            return 2
        model = probe_model(cfg, api_key, out_dir)
        print(f"Selected model: {model}; API key source: {api_key_source}")
        client = QwenClient(cfg, api_key, model)
        label_df = run_batch("label_v2", ids, cases, old_rows, cfg, client, out_dir)
        pred_df = run_batch("prediction_v2", ids, cases, old_rows, cfg, client, out_dir)

    strict_v2, extended_v2 = write_candidate_files(label_df, strict_v1, extended_v1, cfg, out_dir)
    metrics = write_eval_artifacts(pred_df, strict_v2, extended_v2, out_dir)
    comp = write_model_comparison(strict_v2, extended_v2, out_dir)
    write_figures(out_dir, cfg)
    report = write_word_report(out_dir, cfg, model)
    label_success = float((label_df.get("label_api_status", pd.Series(dtype=str)) == "api_available").mean()) if len(label_df) else 0.0
    pred_success = float((pred_df.get("prediction_api_status", pd.Series(dtype=str)) == "api_available").mean()) if len(pred_df) else 0.0
    audit_status = "complete" if min(label_success, pred_success) >= float(cfg["eval"].get("min_success_rate", 0.95)) else "partial_api_failures"
    write_manifest_files(out_dir, cfg_path, cfg, model, api_key_source, audit_status)
    summary = {
        "out_dir": str(out_dir),
        "selected_model": model,
        "label_success_rate": label_success,
        "prediction_success_rate": pred_success,
        "audit_status": audit_status,
        "metrics": metrics,
        "model_comparison_rows": comp.to_dict("records") if len(comp) else [],
        "word_report": str(report) if report else None,
    }
    (out_dir / "quick_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2)[:4000])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
