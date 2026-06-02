from __future__ import annotations

import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
IJPM_DIR = DESKTOP / "IJPM-PDW-2rd"
SOURCE_DOC = IJPM_DIR / "DelayDispute_Copilot_Integrated_SCI_20260416.docx"
SECTION_SOURCE = REPO_ROOT / "paper_assets" / "text" / "paper_consolidation_sections_20260420.md"
LIMITATION_SOURCE = REPO_ROOT / "paper_assets" / "text" / "limitations_statement.md"
VALIDATION_SOURCE = REPO_ROOT / "paper_assets" / "text" / "validation_statement.md"
TABLE_DIR = REPO_ROOT / "paper_assets" / "tables"
FIG_DIR = REPO_ROOT / "paper_assets" / "figures"
OUTPUT_DIR = IJPM_DIR
OUTPUT_NAME = "DelayDispute_Copilot_IEEE_TEM_Consolidated_20260420.docx"
WORKSPACE_COPY = REPO_ROOT / "paper_assets" / "manuscript" / OUTPUT_NAME


def clean_text(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace("  ", " ").strip()
    return text


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def extract_section(markdown: str, heading: str) -> str:
    pattern = re.compile(rf"(?ms)^## {re.escape(heading)}\s*\n(.*?)(?=^## |\Z)")
    match = pattern.search(markdown)
    if not match:
        raise ValueError(f"Missing section: {heading}")
    return match.group(1).strip()


def extract_subsections(markdown_section: str) -> list[tuple[str, list[str]]]:
    pattern = re.compile(r"(?ms)^### (.+?)\s*\n(.*?)(?=^### |\Z)")
    blocks: list[tuple[str, list[str]]] = []
    for title, body in pattern.findall(markdown_section):
        paragraphs = [clean_text(p) for p in re.split(r"\n\s*\n", body.strip()) if p.strip()]
        blocks.append((clean_text(title), paragraphs))
    return blocks


def extract_paragraphs(markdown_section: str) -> list[str]:
    return [clean_text(p) for p in re.split(r"\n\s*\n", markdown_section.strip()) if p.strip()]


def section_title(title: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)?\s*", "", clean_text(title))


def set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.gutter = Cm(0)


def set_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)


def set_run_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_centered_paragraph(doc: Document, text: str, size: float, bold: bool = False, italic: bool = False) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(clean_text(text))
    set_run_font(run, size=size, bold=bold, italic=italic)
    para.space_after = Pt(3)


def add_body_paragraph(doc: Document, text: str, first_line_indent: Cm | None = Cm(0.74)) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = first_line_indent
    run = para.add_run(clean_text(text))
    set_run_font(run)


def add_section_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(clean_text(text))
    set_run_font(run, size=12, bold=True)


def add_subheading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(clean_text(text))
    set_run_font(run, size=10.5, bold=True)


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(clean_text(text))
    set_run_font(run, size=9.5, italic=True)


def add_note(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(clean_text(text))
    set_run_font(run, size=9.0)


def shade_cell(cell, fill: str = "D9D9D9") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(clean_text(str(text)))
    set_run_font(run, size=9.2, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_dataframe_table(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        set_table_text(cell, col, bold=True)
        shade_cell(cell)

    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, value in enumerate(row):
            display = value
            if isinstance(display, float):
                display = f"{display:.3f}"
            set_table_text(table.cell(i, j), display, bold=False)

    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, width_inches: float = 6.2) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    para.paragraph_format.space_after = Pt(2)


def build_table1() -> pd.DataFrame:
    df = pd.read_csv(TABLE_DIR / "table1_candidate_benchmark_profile.csv")
    out = pd.DataFrame(
        {
            "Benchmark": df["benchmark"].str.replace(" benchmark", "", regex=False),
            "Dataset": df["dataset_name"],
            "n": df["n_cases"],
            "Support": df["support_n"].astype(str) + " (" + df["support_ratio"].map(lambda x: f"{x:.3f}") + ")",
            "Partial": df["partial_n"].astype(str) + " (" + df["partial_ratio"].map(lambda x: f"{x:.3f}") + ")",
            "Not support": df["not_support_n"].astype(str) + " (" + df["not_support_ratio"].map(lambda x: f"{x:.3f}") + ")",
            "Owner": df["owner_n"].astype(str) + " (" + df["owner_ratio"].map(lambda x: f"{x:.3f}") + ")",
            "Contractor": df["contractor_n"].astype(str) + " (" + df["contractor_ratio"].map(lambda x: f"{x:.3f}") + ")",
            "Unknown": df["unknown_n"].astype(str) + " (" + df["unknown_ratio"].map(lambda x: f"{x:.3f}") + ")",
        }
    )
    return out


def build_table2() -> pd.DataFrame:
    df = pd.read_csv(TABLE_DIR / "table2_main_outcome_results.csv")
    model_map = {
        "majority_class": "Majority class",
        "rule_baseline": "Rule baseline",
        "tfidf_logreg": "TF-IDF + LogReg",
        "tfidf_linearsvc": "TF-IDF + LinearSVC",
        "tfidf_multinomialnb": "TF-IDF + MultinomialNB",
        "current_hybrid_baseline": "Current hybrid baseline",
        "paesc_hybrid": "PAESC hybrid",
    }
    dataset_map = {
        "candidate_gold_strict_v1": "Strict",
        "candidate_gold_extended_v1": "Extended",
    }
    out = df.copy()
    out["Dataset"] = out["dataset_name"].map(dataset_map)
    out["Model"] = out["model_name"].map(model_map)
    out["95% CI (Macro-F1)"] = out["macro_f1_ci_low"].map(lambda x: f"{x:.3f}") + "–" + out["macro_f1_ci_high"].map(lambda x: f"{x:.3f}")
    out = out[["Dataset", "Model", "accuracy", "macro_f1", "weighted_f1", "95% CI (Macro-F1)"]]
    out.columns = ["Dataset", "Model", "Accuracy", "Macro-F1", "Weighted-F1", "95% CI (Macro-F1)"]
    return out


def build_table3() -> pd.DataFrame:
    df = pd.read_csv(TABLE_DIR / "table3_ablation_leakage_stress.csv")
    dataset_map = {
        "candidate_gold_strict_v1": "Strict",
        "candidate_gold_extended_v1": "Extended",
    }
    setting_map = {
        "full_model": "Full model",
        "remove_pre_decision_constraint": "Remove pre-decision constraint",
        "remove_structured_events": "Remove structured events",
        "remove_procedural_signals": "Remove procedural signals",
        "remove_evidence_chain": "Remove evidence chain",
        "remove_responsibility_head": "Remove responsibility head",
        "remove_retrieval": "Remove retrieval",
        "remove_irac_verifier": "Remove verifier",
    }
    out = df.copy()
    out["Dataset"] = out["dataset_name"].map(dataset_map)
    out["Setting"] = out["ablation_setting"].map(setting_map)
    out = out[
        [
            "Dataset",
            "Setting",
            "accuracy",
            "macro_f1",
            "responsibility_macro_f1",
            "valid_span_rate",
            "pre_decision_span_rate",
            "high_dispute_rate",
        ]
    ]
    out.columns = [
        "Dataset",
        "Setting",
        "Outcome Acc.",
        "Outcome Macro-F1",
        "Resp. Macro-F1",
        "Valid span",
        "Pre-decision span",
        "High-dispute",
    ]
    return out


def build_table4() -> pd.DataFrame:
    df = pd.read_csv(TABLE_DIR / "table4_responsibility_auditability_summary.csv")
    out = df[
        [
            "benchmark",
            "responsibility_accuracy",
            "responsibility_macro_f1",
            "candidate_unknown_ratio",
            "strongest_class",
            "strongest_class_f1",
            "weakest_class",
            "weakest_class_f1",
            "valid_span_rate",
            "pre_decision_span_rate",
            "duplicate_chain_rate",
            "role_coverage_rate",
        ]
    ].copy()
    out.columns = [
        "Benchmark",
        "Resp. Acc.",
        "Resp. Macro-F1",
        "Unknown ratio",
        "Strongest class",
        "Strongest F1",
        "Weakest class",
        "Weakest F1",
        "Valid span",
        "Pre-decision span",
        "Duplicate chain",
        "Role coverage",
    ]
    return out


def build_table5() -> pd.DataFrame:
    df = pd.read_csv(TABLE_DIR / "table5_forensic_claim_boundary.csv")
    out = df[
        [
            "result_group",
            "run_name",
            "dataset_name",
            "model_name",
            "accuracy",
            "macro_f1",
            "audit_status",
            "claim_tier",
        ]
    ].copy()
    out.columns = [
        "Result group",
        "Run",
        "Dataset",
        "Model",
        "Accuracy",
        "Macro-F1",
        "Audit status",
        "Claim tier",
    ]
    return out


def get_source_doc() -> Document:
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"Missing source Word file: {SOURCE_DOC}")
    return Document(str(SOURCE_DOC))


def get_front_matter(source_doc: Document) -> tuple[str, str, str, str]:
    title = clean_text(source_doc.paragraphs[0].text)
    author = clean_text(source_doc.paragraphs[1].text)
    affiliation = clean_text(source_doc.paragraphs[2].text)
    index_terms = clean_text(source_doc.paragraphs[5].text)
    if "—" in index_terms:
        index_terms = index_terms.split("—", 1)[1].strip()
    return title, author, affiliation, index_terms


def get_literature_blocks(source_doc: Document) -> list[tuple[str, list[str]]]:
    return [
        (
            "A. Construction Delay Disputes and Responsibility Determination",
            [clean_text(source_doc.paragraphs[i].text) for i in range(13, 17) if source_doc.paragraphs[i].text.strip()],
        ),
        (
            "B. NLP for Construction Claims and Dispute Analytics",
            [clean_text(source_doc.paragraphs[i].text) for i in range(18, 20) if source_doc.paragraphs[i].text.strip()],
        ),
        (
            "C. Research Gap and Positioning",
            [clean_text(source_doc.paragraphs[i].text) for i in range(21, 25) if source_doc.paragraphs[i].text.strip()],
        ),
    ]


def get_references(source_doc: Document) -> list[str]:
    return [clean_text(p.text) for p in source_doc.paragraphs[87:] if p.text.strip()]


def add_markdown_subsections(doc: Document, subsections: list[tuple[str, list[str]]]) -> None:
    letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    for idx, (title, paragraphs) in enumerate(subsections):
        add_subheading(doc, f"{letters[idx]}. {section_title(title)}")
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)


def add_data_section(doc: Document) -> None:
    add_section_heading(doc, "III. DATA, PROBLEM DEFINITION, AND EVALUATION DESIGN")
    add_subheading(doc, "A. Dataset and Candidate Benchmarks")
    add_body_paragraph(
        doc,
        (
            "The study uses a machine-assisted corpus of approximately 4,592 construction delay dispute cases together "
            "with two candidate evaluation benchmarks: candidate_gold_strict_v1 (250 cases) and "
            "candidate_gold_extended_v1 (500 cases). These datasets are treated as candidate benchmarks rather than "
            "true human-gold test sets, because their labels remain machine-assisted and audit-oriented."
        ),
    )
    add_body_paragraph(
        doc,
        (
            "The central methodological constraint is that all predictive inputs are limited to pre-decision information. "
            "Post-decision text may be used only for label anchoring, candidate-benchmark construction, and forensic audit "
            "stress tests. This separation ensures that the evaluation remains aligned with prospective managerial use "
            "rather than retrospective score maximization."
        ),
    )
    add_caption(doc, "Table 1. Profile of the candidate benchmarks used in the audit-ready evaluation.")
    add_dataframe_table(doc, build_table1())
    add_note(
        doc,
        "Note: Candidate benchmarks are machine-assisted and should not be described as true human gold labels.",
    )
    add_figure(doc, FIG_DIR / "fig2_leakage_benchmark_governance.png", width_inches=6.2)
    add_caption(doc, "Fig. 2. Leakage-control design and candidate-benchmark governance used in the present study.")
    add_subheading(doc, "B. Leakage-Aware Problem Formulation and Evaluation Logic")
    add_body_paragraph(
        doc,
        (
            "The main empirical basis of the paper is the audit-ready run final_eval_20260409_194025, and the claim "
            "boundary is defined by the forensic audit in forensic_audit_20260413_111946. This evidence hierarchy is "
            "explicit: current candidate-benchmark results are the main manuscript basis, whereas older high-scoring runs "
            "are retained only as historical references when they fail the current manifest, recomputation, or leakage-sentinel standard."
        ),
    )
    add_body_paragraph(
        doc,
        (
            "Evaluation therefore proceeds under a strict claim discipline. Outcome prediction is reported with accuracy, "
            "macro-F1, weighted-F1, per-class results, and confusion matrices. Responsibility diagnosis is reported with "
            "accuracy, macro-F1, and audit-oriented uncertainty indicators. Evidence-chain quality is reported separately "
            "through span-validity and pre-decision traceability metrics, rather than being conflated with prediction quality."
        ),
    )


def add_results_section(doc: Document, results_subsections: list[tuple[str, list[str]]]) -> None:
    add_section_heading(doc, "V. RESULTS")
    for idx, (title, paragraphs) in enumerate(results_subsections):
        label = "ABCDE"[idx]
        add_subheading(doc, f"{label}. {section_title(title)}")
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)
        if idx == 1:
            add_caption(doc, "Table 2. Main outcome-prediction results on the strict and extended candidate benchmarks.")
            add_dataframe_table(doc, build_table2())
            add_note(
                doc,
                "Note: The reproducible hybrid baseline is the strongest predictive comparator, whereas PAESC is retained as the integrated audit-ready framework.",
            )
            add_figure(doc, FIG_DIR / "fig3_outcome_prediction_comparison.png", width_inches=6.1)
            add_caption(doc, "Fig. 3. Outcome-prediction performance on the strict and extended candidate benchmarks.")
            add_caption(doc, "Table 3. Ablation and leakage-stress results for the audit-ready framework.")
            add_dataframe_table(doc, build_table3())
            add_note(
                doc,
                "Note: Rows removing the pre-decision constraint are stress tests only and are not deployment-valid settings.",
            )
        elif idx == 2:
            add_caption(doc, "Table 4. Summary of responsibility-diagnosis performance and evidence-chain auditability.")
            add_dataframe_table(doc, build_table4())
            add_note(
                doc,
                "Note: Responsibility-diagnosis and evidence-chain outputs are machine-evaluated and audit-ready rather than fully human-validated.",
            )
            add_figure(doc, FIG_DIR / "fig4_responsibility_auditability.png", width_inches=6.1)
            add_caption(doc, "Fig. 4. Responsibility-diagnosis performance and evidence-chain auditability on the candidate benchmarks.")
        elif idx == 3:
            add_caption(doc, "Table 5. Forensic claim boundary for historical and current results.")
            add_dataframe_table(doc, build_table5())
            add_note(
                doc,
                "Note: Tier B results are claimable with caution; Tier C results are historical references only.",
            )
            add_figure(doc, FIG_DIR / "fig5_forensic_claim_boundary.png", width_inches=6.1)
            add_caption(doc, "Fig. 5. Forensic audit of leakage-related score inflation risk and result claim boundary.")
        elif idx == 4:
            add_figure(doc, FIG_DIR / "fig6_representative_case_panel.png", width_inches=6.2)
            add_caption(
                doc,
                "Fig. 6. Representative audit-ready case panel showing pre-decision evidence, predicted outcome, structured responsibility diagnosis, and uncertainty cues.",
            )


def build_document() -> Path:
    source_doc = get_source_doc()
    title, author, affiliation, index_terms = get_front_matter(source_doc)
    references = get_references(source_doc)
    literature_blocks = get_literature_blocks(source_doc)

    markdown = read_text(SECTION_SOURCE)
    abstract_paragraphs = extract_paragraphs(extract_section(markdown, "Abstract"))
    managerial_paragraphs = extract_paragraphs(extract_section(markdown, "Managerial Relevance Statement"))
    intro_paragraphs = extract_paragraphs(extract_section(markdown, "1. Introduction"))
    method_subsections = extract_subsections(extract_section(markdown, "2. Method"))
    results_subsections = extract_subsections(extract_section(markdown, "3. Results"))
    discussion_paragraphs = extract_paragraphs(extract_section(markdown, "4. Discussion"))
    conclusion_paragraphs = extract_paragraphs(extract_section(markdown, "5. Conclusion"))
    validation_paragraphs = extract_paragraphs(read_text(VALIDATION_SOURCE))
    limitation_paragraphs = extract_paragraphs(read_text(LIMITATION_SOURCE))

    doc = Document()
    set_page_layout(doc)
    set_document_styles(doc)

    add_centered_paragraph(doc, title, size=15, bold=True)
    add_centered_paragraph(doc, author, size=11)
    add_centered_paragraph(doc, affiliation, size=10)
    doc.add_paragraph()

    add_body_paragraph(doc, f"Abstract—{abstract_paragraphs[0]}", first_line_indent=None)
    for paragraph in abstract_paragraphs[1:]:
        add_body_paragraph(doc, paragraph, first_line_indent=None)

    add_body_paragraph(doc, f"Managerial Relevance Statement—{managerial_paragraphs[0]}", first_line_indent=None)
    for paragraph in managerial_paragraphs[1:]:
        add_body_paragraph(doc, paragraph, first_line_indent=None)

    add_body_paragraph(doc, f"Index Terms—{index_terms}", first_line_indent=None)

    add_section_heading(doc, "I. INTRODUCTION")
    for paragraph in intro_paragraphs:
        add_body_paragraph(doc, paragraph)

    add_section_heading(doc, "II. LITERATURE REVIEW")
    for heading, paragraphs in literature_blocks:
        add_subheading(doc, heading)
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)

    add_data_section(doc)

    add_section_heading(doc, "IV. METHODOLOGY")
    add_figure(doc, FIG_DIR / "fig1_delaydispute_workflow.png", width_inches=6.1)
    add_caption(doc, "Fig. 1. Leakage-aware workflow of DelayDispute Copilot for construction delay disputes.")
    add_markdown_subsections(doc, method_subsections)

    add_results_section(doc, results_subsections)

    add_section_heading(doc, "VI. DISCUSSION")
    for paragraph in discussion_paragraphs:
        add_body_paragraph(doc, paragraph)
    add_subheading(doc, "A. Validation Boundary and Limitations")
    for paragraph in validation_paragraphs + limitation_paragraphs:
        add_body_paragraph(doc, paragraph)

    add_section_heading(doc, "VII. CONCLUSION")
    for paragraph in conclusion_paragraphs:
        add_body_paragraph(doc, paragraph)

    add_subheading(doc, "Supplementary Material Note")
    add_body_paragraph(
        doc,
        (
            "Detailed per-class metrics, leakage-sentinel tables, forensic delta audits, expanded ablations, and additional "
            "representative cases are preserved in the repository supplementary package under paper_assets and results and "
            "should be cited as supplementary evidence rather than promoted into the main claim set."
        ),
    )

    add_section_heading(doc, "REFERENCES")
    for ref in references:
        add_body_paragraph(doc, ref, first_line_indent=None)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / OUTPUT_NAME
    doc.save(str(output_path))

    WORKSPACE_COPY.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(output_path, WORKSPACE_COPY)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(path)
